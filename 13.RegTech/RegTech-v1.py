from __future__ import annotations # MUST be the first import
import streamlit as st
import os
import requests
from typing import List, Dict, Generator, Optional, Union
import json
import datetime

# --- Library Imports ---
# We import everything here to ensure names are defined.
try:
    from langchain_core.documents import Document
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
    LIBRARIES_AVAILABLE = True
except ImportError as e:
    st.error(f"A required library is not installed. Please check the installation instructions in the original script. Error: {e}")
    LIBRARIES_AVAILABLE = False


# --- Configuration ---
# Set your folder path and DeepSeek API Key here.
# The folder should be in the same directory as this script.
FOLDER_PATH = "./legal_docs"
API_KEY = "sk-cfbf269ccbbc4e94aa69df94c2a25739" # Your DeepSeek API Key
RULES_FILE_NAME = "RULES.txt" # The name of the mandatory rules file.

# --- Tesseract Installation Check ---
# Pytesseract requires Google's Tesseract-OCR engine. If you get an error,
# follow the installation instructions from the original script and uncomment the line below.
# pytesseract.pytesseract.tesseract_cmd = r'<path_to_your_tesseract_executable>'


# --- DeepSeek API Configuration ---
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

# --- Helper Functions ---

def read_file_content(file: Union[str, st.runtime.uploaded_file_manager.UploadedFile]) -> tuple[str, str]:
    """
    Reads content from a file path or a Streamlit UploadedFile object.
    Supports .pdf, .docx, .txt, and image files.

    Args:
        file: A file path (str) or a Streamlit UploadedFile object.

    Returns:
        A tuple containing the file name and its extracted text content.
    """
    text = ""
    file_name = ""
    try:
        if isinstance(file, str): # If it's a file path
            file_name = os.path.basename(file)
            file_path = file
            file_opener = open(file_path, 'rb')
        else: # If it's a Streamlit UploadedFile
            file_name = file.name
            file_opener = file

        # Handle .pdf files
        if file_name.endswith(".pdf"):
            reader = pypdf.PdfReader(file_opener)
            for page in reader.pages:
                text += page.extract_text() or ""
        # Handle .docx files
        elif file_name.endswith(".docx"):
            doc = docx.Document(file_opener)
            for para in doc.paragraphs:
                text += para.text + "\n"
        # Handle .txt files
        elif file_name.endswith(".txt"):
            # We need to handle the UploadedFile differently for text
            if not isinstance(file, str):
                text = file.getvalue().decode('utf-8')
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    text = f.read()
        # Handle image files (.png, .jpg, .jpeg) using OCR
        elif file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
            try:
                text = pytesseract.image_to_string(Image.open(file_opener))
            except pytesseract.TesseractNotFoundError:
                st.error("Tesseract is not installed or not in your PATH. Please check instructions.")
            except Exception as ocr_e:
                st.error(f"OCR failed for {file_name}: {ocr_e}")

        if isinstance(file, str): # Close the file if we opened it
            file_opener.close()

    except Exception as e:
        st.error(f"Failed to read or process {file_name}: {e}")

    return file_name, text


def load_documents_from_folder(folder_path: str) -> List[Document]:
    """
    Loads all supported files from a folder into a list of Document objects.
    """
    documents = []
    if not os.path.isdir(folder_path):
        st.error(f"The provided path '{folder_path}' is not a valid directory.")
        return []

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        _, text = read_file_content(file_path)
        if text:
            documents.append(Document(page_content=text, metadata={"source": file_name}))
    return documents


def get_deepseek_response(prompt: str, api_key: str, chat_history: List[Dict[str, str]]) -> Generator[str, None, None]:
    """
    Sends a request to the DeepSeek API and streams the response.
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    # The prompt is already fully formatted, so we just add it to the history.
    messages = chat_history + [{"role": "user", "content": prompt}]
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "stream": True
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, stream=True)
        response.raise_for_status()
        for chunk in response.iter_lines():
            if chunk:
                decoded_chunk = chunk.decode('utf-8')
                if decoded_chunk.startswith('data: '):
                    try:
                        json_data = json.loads(decoded_chunk[6:])
                        if 'choices' in json_data and json_data['choices']:
                            content = json_data['choices'][0].get('delta', {}).get('content', '')
                            if content:
                                yield content
                    except json.JSONDecodeError:
                        continue # Ignore malformed chunks
    except requests.exceptions.RequestException as e:
        error_message = f"API request failed: {e}"
        st.error(error_message)
        yield error_message
    except Exception as e:
        error_message = f"An unexpected error occurred during API call: {e}"
        st.error(error_message)
        yield error_message

# --- Custom CSS for Enhanced Design ---
def apply_custom_css():
    """Applies custom CSS to improve the application's design."""
    custom_css = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
        
        html, body, [class*="st-"] {
            font-family: 'Inter', sans-serif;
        }

        /* Main app background */
        .stApp {
            background-color: #111827; /* Dark blue-gray background */
            color: #F9FAFB;
        }

        /* Hide Streamlit's default header and footer */
        #MainMenu {display: none;}
        footer {display: none;}
        .stApp > header {
            background-color: transparent;
        }

        /* Chat message styling */
        [data-testid="chat-message-container"] {
            border-radius: 12px;
            padding: 16px;
            margin-bottom: 12px;
            max-width: 85%;
        }

        /* Assistant's message bubble */
        [data-testid="chat-message-container"]:has(div[data-testid="stChatMessageContent"][aria-label="assistant's message"]) {
            background-color: #374151; /* Gray-700 */
        }
        
        /* User's message bubble */
        [data-testid="chat-message-container"]:has(div[data-testid="stChatMessageContent"][aria-label="user's message"]) {
            background-color: #1E40AF; /* Blue-800 */
        }
        
        /* Style the chat input box */
        [data-testid="stChatInput"] {
            background-color: #1F2937; /* Gray-800 */
            border-radius: 12px;
            border: 1px solid #4B5563; /* Gray-600 */
        }
        [data-testid="stChatInput"] textarea {
            color: #F9FAFB;
        }

        /* Sidebar styling */
        [data-testid="stSidebar"] {
            background-color: #1F2937; /* Gray-800 */
            border-right: 1px solid #374151;
        }

        /* Custom styling for document selection */
        .document-selector-container {
            background-color: #374151;
            border-radius: 8px;
            padding: 10px;
            margin-bottom: 1rem;
        }
        .document-selector-container .stCheckbox {
            border-bottom: 1px solid #4B5563;
            padding: 5px 0;
        }
        .document-selector-container .stCheckbox:last-child {
            border-bottom: none;
        }
    </style>
    """
    st.markdown(custom_css, unsafe_allow_html=True)

def build_prompt(task_description: str, rules: str, context: str, question: Optional[str] = None) -> str:
    """Constructs a complete prompt for the LLM based on the task."""
    
    question_section = f"""
    USER'S QUESTION/REQUEST:
    ---
    {question}
    ---
    """ if question else ""

    return f"""
    You are a highly specialized legal assistant, "Mofti Pro+". You must strictly follow the rules provided below in all your responses.

    MANDATORY RULES:
    ---
    {rules}
    ---

    Your task is to perform the following action: {task_description}
    Base your response *only* on the provided context documents. Do not use any external knowledge.
    If the information is not found in the context, state that clearly.
    When answering questions or summarizing, you MUST cite the source file(s) you used from the context.

    CONTEXT DOCUMENTS:
    ---
    {context}
    ---
    {question_section}
    """

# --- Streamlit App UI ---
def main():
    """The main function that runs the Streamlit application."""
    st.set_page_config(
        page_title="Mofti Pro+", 
        page_icon="⚖️",
        layout="wide",
        initial_sidebar_state="expanded" 
    )
    
    apply_custom_css()

    st.title("⚖️ Mofti Pro+")
    
    if not LIBRARIES_AVAILABLE:
        st.stop()
    if not API_KEY or "sk-" not in API_KEY:
        st.error("Your DeepSeek API Key is not set correctly at the top of the script.")
        st.stop()

    # --- Initialize Session State ---
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "loaded_docs" not in st.session_state:
        st.session_state.loaded_docs = []
    if "docs_loaded_first_time" not in st.session_state:
        st.session_state.docs_loaded_first_time = False
    if "temp_file_content" not in st.session_state:
        st.session_state.temp_file_content = None
    if "selected_docs_map" not in st.session_state:
        st.session_state.selected_docs_map = {}


    # --- Document Loading (runs only once or when refreshed) ---
    if not st.session_state.docs_loaded_first_time:
        with st.spinner("Loading documents from permanent storage..."):
            if not os.path.exists(FOLDER_PATH):
                os.makedirs(FOLDER_PATH)
            st.session_state.loaded_docs = load_documents_from_folder(FOLDER_PATH)
            
            # Initialize the selection map
            all_doc_names = [doc.metadata['source'] for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME]
            st.session_state.selected_docs_map = {name: True for name in all_doc_names}

            st.session_state.docs_loaded_first_time = True
            if st.session_state.loaded_docs:
                st.toast(f"Successfully loaded {len(st.session_state.loaded_docs)} documents.", icon="✅")
            else:
                st.warning("No permanent documents found. You can upload them in the sidebar.")

    # --- Sidebar ---
    with st.sidebar:
        st.header("File Management")

        # Feature: Permanent File Uploader
        uploaded_files = st.file_uploader(
            "Upload permanent documents",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="Files uploaded here are saved to the server and will be available every time you open the app."
        )
        if uploaded_files:
            with st.spinner("Saving and processing permanent files..."):
                for uploaded_file in uploaded_files:
                    with open(os.path.join(FOLDER_PATH, uploaded_file.name), "wb") as f:
                        f.write(uploaded_file.getbuffer())
            st.toast(f"Uploaded {len(uploaded_files)} files! Reloading...", icon="�")
            st.session_state.docs_loaded_first_time = False
            st.rerun()

        st.markdown("---")
        
        selectable_docs = [doc for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME]
        
        if selectable_docs:
            st.header("Select Documents for Context")
            st.markdown('<div class="document-selector-container">', unsafe_allow_html=True)
            for doc in selectable_docs:
                doc_name = doc.metadata['source']
                is_selected = st.checkbox(doc_name, value=st.session_state.selected_docs_map.get(doc_name, True), key=f"cb_{doc_name}")
                st.session_state.selected_docs_map[doc_name] = is_selected
            st.markdown('</div>', unsafe_allow_html=True)

            st.session_state.selected_docs = [
                doc for doc in selectable_docs if st.session_state.selected_docs_map.get(doc.metadata['source'])
            ]
        else:
            st.info("No documents to select. Please upload some files.")
            st.session_state.selected_docs = []

        st.markdown("---")
        st.header("Actions")

        # --- New Feature Buttons ---
        if st.button("📄 Summarize a Single File", use_container_width=True):
            st.session_state.action = "summarize_single"
            st.rerun()

        if st.button("🔄 Compare Two Documents", use_container_width=True):
            st.session_state.action = "compare_docs"
            st.rerun()
        
        if st.button("🔍 Check for Contradictions", use_container_width=True):
            st.session_state.action = "check_contradictions"
            st.rerun()
        
        if st.button("📖 Define a Legal Term", use_container_width=True):
            st.session_state.action = "define_term"
            st.rerun()

        if st.button("✍️ Draft a Legal Clause", use_container_width=True):
            st.session_state.action = "draft_clause"
            st.rerun()

        st.markdown("---")
        
        if st.button("🗑️ Clear & Start New Topic", use_container_width=True):
            st.session_state.messages = []
            st.session_state.action = None
            st.toast("Chat history cleared!", icon="🧹")
            st.rerun()


    # --- Main Chat Interface ---
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # --- Handle Actions ---
    if "action" in st.session_state and st.session_state.action:
        action = st.session_state.action
        
        # This container will hold the UI for the specific action
        action_placeholder = st.container()

        # Shared logic for context and rules
        rules_doc = next((doc for doc in st.session_state.loaded_docs if doc.metadata['source'] == RULES_FILE_NAME), None)
        rules_content = rules_doc.page_content if rules_doc else "No specific rules provided."
        
        prompt_to_send = None
        user_message_content = ""
        context_text = ""

        with action_placeholder:
            if action == "summarize_single":
                st.info("Select a file to summarize.")
                file_to_summarize = st.selectbox("Choose a document:", [d.metadata['source'] for d in selectable_docs])
                if st.button("Generate Summary"):
                    doc_to_summarize = next((d for d in selectable_docs if d.metadata['source'] == file_to_summarize), None)
                    if doc_to_summarize:
                        context_text = f"Source: {doc_to_summarize.metadata['source']}\n\nContent: {doc_to_summarize.page_content}"
                        task = f"Provide a concise summary of the key points, findings, and conclusions from the document '{file_to_summarize}'."
                        prompt_to_send = build_prompt(task, rules_content, context_text)
                        user_message_content = f"Please summarize the document: {file_to_summarize}"

            elif action == "compare_docs":
                st.info("Select two files to compare.")
                doc_names = [d.metadata['source'] for d in selectable_docs]
                col1, col2 = st.columns(2)
                file1 = col1.selectbox("First document:", doc_names, key="comp1")
                file2 = col2.selectbox("Second document:", doc_names, key="comp2", index=min(1, len(doc_names)-1))
                if st.button("Run Comparison"):
                    if file1 == file2:
                        st.warning("Please select two different documents.")
                    else:
                        doc1 = next((d for d in selectable_docs if d.metadata['source'] == file1), None)
                        doc2 = next((d for d in selectable_docs if d.metadata['source'] == file2), None)
                        context_text = f"Document 1 Source: {doc1.metadata['source']}\nContent:\n{doc1.page_content}\n\n---\n\nDocument 2 Source: {doc2.metadata['source']}\nContent:\n{doc2.page_content}"
                        task = f"Compare and contrast the two provided documents ('{file1}' and '{file2}'). Highlight key similarities, differences, and any potential conflicts between them."
                        prompt_to_send = build_prompt(task, rules_content, context_text)
                        user_message_content = f"Please compare '{file1}' and '{file2}'."

            elif action == "check_contradictions":
                if not st.session_state.selected_docs:
                     st.warning("Please select at least one document from the sidebar to check for contradictions.")
                else:
                    st.info("Click the button to check for contradictions in the selected documents.")
                    if st.button("Find Contradictions"):
                        context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in st.session_state.selected_docs])
                        task = "Analyze the provided documents and identify any potential contradictions, inconsistencies, or conflicting statements within or between them. List each potential contradiction and cite the source documents."
                        prompt_to_send = build_prompt(task, rules_content, context_text)
                        user_message_content = f"Check for contradictions in the following documents: {', '.join([d.metadata['source'] for d in st.session_state.selected_docs])}"

            elif action == "define_term":
                st.info("Enter a legal term to define based on the selected documents.")
                term = st.text_input("Term to Define:")
                if st.button("Get Definition"):
                    if not term:
                        st.warning("Please enter a term.")
                    elif not st.session_state.selected_docs:
                        st.warning("Please select at least one document to provide context.")
                    else:
                        context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in st.session_state.selected_docs])
                        task = f"Define the legal term '{term}' using only the information available in the provided context documents. If the term is not found, state that."
                        prompt_to_send = build_prompt(task, rules_content, context_text, question=f"Define: {term}")
                        user_message_content = f"Define the term: '{term}'"

            elif action == "draft_clause":
                st.info("Describe the legal clause you want to draft.")
                description = st.text_area("Clause Description:")
                if st.button("Generate Draft"):
                    if not description:
                        st.warning("Please describe the clause.")
                    elif not st.session_state.selected_docs:
                        st.warning("Please select at least one document to provide context for style and terminology.")
                    else:
                        context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in st.session_state.selected_docs])
                        task = f"Draft a legal clause based on the following description: '{description}'. Use the style, terminology, and legal framework found in the provided context documents."
                        prompt_to_send = build_prompt(task, rules_content, context_text, question=f"Draft a clause for: {description}")
                        user_message_content = f"Draft a legal clause for me. Description: {description}"

        if prompt_to_send and user_message_content:
            st.session_state.messages.append({"role": "user", "content": user_message_content})
            with st.chat_message("assistant"):
                with st.spinner("The legal assistant is working..."):
                    response_placeholder = st.empty()
                    full_response = ""
                    chat_history_for_api = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1]]
                    for chunk in get_deepseek_response(prompt_to_send, API_KEY, chat_history_for_api):
                        full_response += chunk
                        response_placeholder.markdown(full_response + "▌")
                    response_placeholder.markdown(full_response)
            st.session_state.messages.append({"role": "assistant", "content": full_response})
            st.session_state.action = None # Clear action after completion
            st.rerun()

    # --- Temporary File Uploader & Chat Input ---
    # This section is now at the bottom of the main area
    st.markdown("---")
    temp_file = st.file_uploader(
        "Upload a temporary file for your next question",
        type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
        help="This file will be used for the next question only and then discarded."
    )

    if temp_file:
        with st.spinner("Reading temporary file..."):
            file_name, text = read_file_content(temp_file)
            if text:
                st.session_state.temp_file_content = Document(page_content=text, metadata={"source": f"Temporary file: {file_name}"})
                st.info(f"Ready to use temporary file: `{file_name}`. Ask your question below.")
            else:
                st.warning(f"Could not read content from `{file_name}`.")
                st.session_state.temp_file_content = None

    if prompt := st.chat_input("Ask a question about your documents..."):
        # An action from the sidebar might be active, we should clear it if the user starts typing
        if st.session_state.get("action"):
            st.session_state.action = None

        all_context_docs = st.session_state.get("selected_docs", [])
        if st.session_state.get("temp_file_content"):
            all_context_docs = all_context_docs + [st.session_state.temp_file_content]

        if not all_context_docs:
            st.warning("Please select at least one document or upload a temporary file to ask a question.")
            st.stop()
        
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("assistant"):
            with st.spinner("The legal assistant is thinking..."):
                # Load mandatory rules
                rules_doc = next((doc for doc in st.session_state.loaded_docs if doc.metadata['source'] == RULES_FILE_NAME), None)
                rules_content = rules_doc.page_content if rules_doc else "No specific rules provided. Follow general instructions."
                
                # Build context from all relevant docs (permanent + temporary)
                context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in all_context_docs])
                
                task = "Answer the user's question based on the provided context."
                prompt_to_send = build_prompt(task, rules_content, context_text, prompt)

                response_placeholder = st.empty()
                full_response = ""
                
                chat_history_for_api = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1]]

                for chunk in get_deepseek_response(prompt_to_send, API_KEY, chat_history_for_api):
                    full_response += chunk
                    response_placeholder.markdown(full_response + "▌")
                response_placeholder.markdown(full_response)

        st.session_state.messages.append({"role": "assistant", "content": full_response})
        # Clear the temporary file after it has been used
        st.session_state.temp_file_content = None
        st.rerun()


if __name__ == "__main__":
    main()