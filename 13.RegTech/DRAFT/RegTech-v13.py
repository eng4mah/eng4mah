
#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
RegTech Assistant – Streamlit app
---------------------------------

The same code works on Streamlit Cloud (GitHub‑mounted) **and** on a local
machine.  No modifications are required to switch between the two
environments.

Prerequisites
~~~~~~~~~~~~~
pip install -r requirements.txt          # streamlit, langchain, pypdf,
                                         # python-docx, pillow, pytesseract,
                                         # requests
# Optional OCR support:
#   Linux:   sudo apt-get install tesseract-ocr
#   macOS:   brew install tesseract

Set your DeepSeek API key as an environment variable
`DEEPSEEK_API_KEY` **or** in `.streamlit/secrets.toml`:

    DEEPSEEK_API_KEY = "sk-..."

Run:
    streamlit run app.py
"""

# ──────────────────────────────────────────────────────────────────────────────
# Imports
# ──────────────────────────────────────────────────────────────────────────────
from __future__ import annotations

import os
import re
import json
import time
import datetime
import requests
from typing import List, Dict, Generator, Optional, Tuple

import streamlit as st

# ──────────────────────────────────────
# Optional heavy libraries (PDF, DOCX, OCR)
# ──────────────────────────────────────
try:
    from langchain_core.documents import Document
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
    LIBRARIES_AVAILABLE = True
except Exception as e:                     # catches ImportError and any runtime errors
    st.error(f"❗ مكتبة ناقصة أو غير مُثبتة: {e}")
    LIBRARIES_AVAILABLE = False

# ──────────────────────────────────────
# Configuration – paths are resolved dynamically
# ──────────────────────────────────────
def get_repo_root() -> str:
    """
    Detect the repository root.

    1️⃣ Cloud layout: the repo is mounted under /mount/src/<repo_name>/.
        Walk upwards until a folder containing `13.RegTech` is found.

    2️⃣ Fallback: assume the script is running locally and return the
        directory that contains this file.
    """
    cwd = os.path.abspath(os.path.dirname(__file__))
    candidate = cwd
    for _ in range(5):
        if os.path.isdir(os.path.join(candidate, "13.RegTech")):
            return candidate
        candidate = os.path.abspath(os.path.join(candidate, ".."))
    return cwd   # local development fallback

# Resolve the root once – used everywhere else
REPO_ROOT = get_repo_root()

# Fixed sub‑folders (relative to the detected root)
FOLDER_PATH = os.path.join(REPO_ROOT, "13.RegTech", "legal_docs")
SUMMARY_FOLDER_PATH = os.path.join(REPO_ROOT, "13.RegTech", "summary")
RULES_FILE_NAME = "RULES.txt"
RULES_FILE_PATH = os.path.join(REPO_ROOT, "13.RegTech", RULES_FILE_NAME)

# API configuration
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"
API_KEY = os.getenv("DEEPSEEK_API_KEY") or st.secrets.get("DEEPSEEK_API_KEY", "")

# Feature flags ---------------------------------------------------------------
USE_AUTO_SUMMARIZATION = True
USE_SUMMARY_FOR_SELECTION = True
USE_TWO_STEP_SELECTION = True
SELECTION_MAX_RETRIES = 3
MAX_DOCS_FOR_SELECTION = 1   # strict limit as requested

# ──────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────
def safe_run(fn):
    """Wrap the whole Streamlit app so uncaught exceptions are shown, not fatal."""
    def wrapper(*args, **kwargs):
        try:
            return fn(*args, **kwargs)
        except Exception as exc:
            st.error(f"❗ حدث خطأ غير متوقع: {exc}")
            st.exception(exc)
    return wrapper


def read_file_content(file_path: str) -> Tuple[str, str]:
    """Read text from PDF, DOCX, TXT, or image (OCR). Returns (file_name, text)."""
    text = ""
    file_name = os.path.basename(file_path)

    try:
        with open(file_path, "rb") as f:
            if file_name.lower().endswith(".pdf"):
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            elif file_name.lower().endswith(".docx"):
                doc = docx.Document(f)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif file_name.lower().endswith((".png", ".jpg", ".jpeg")):
                try:
                    text = pytesseract.image_to_string(Image.open(f))
                except pytesseract.TesseractNotFoundError:
                    st.error("❗ Tesseract غير مُثبت على الجهاز.")
                except Exception as ocr_e:
                    st.error(f"❗ فشل قراءة الصورة {file_name}: {ocr_e}")
            else:   # txt or any other plain‑text fallback
                with open(file_path, "r", encoding="utf-8") as txt_f:
                    text = txt_f.read()
    except Exception as e:
        st.error(f"❗ خطأ قراءة الملف {file_name}: {e}")

    return file_name, text


def get_document_names_from_folder(folder_path: str) -> List[str]:
    """Return a list of supported document names inside *folder_path*."""
    if not os.path.isdir(folder_path):
        st.warning(f"⚠️ المجلد غير موجود: {folder_path}")
        return []

    supported = [
        f for f in os.listdir(folder_path)
        if f != RULES_FILE_NAME and f.lower().endswith(
            (".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg")
        )
    ]
    return supported


def load_rules_file(rules_path: str) -> str:
    """Load the RULES.txt file (or return a default message)."""
    if os.path.exists(rules_path):
        _, content = read_file_content(rules_path)
        return content
    return "لا توجد تعليمات خاصة. اتبع التعليمات العامة."


def get_deepseek_response_blocking(prompt: str, api_key: str) -> Optional[str]:
    """Call DeepSeek once (non‑streaming) and return the answer."""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "stream": False,
    }

    try:
        response = requests.post(
            DEEPSEEK_API_URL,
            headers=headers,
            json=payload,
            timeout=90,
        )
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.RequestException as e:
        st.error(f"❗ فشل الاتصال بالـ API: {e}")
    except Exception as e:
        st.error(f"❗ خطأ غير متوقع عند طلب الاستجابة: {e}")
    return None


def generate_and_save_summary(
    file_name: str,
    doc_folder: str,
    summary_folder: str,
    api_key: str,
) -> None:
    """Create a one‑paragraph summary for *file_name* if it does not exist."""
    summary_file = f"{os.path.splitext(file_name)[0]}_summary.txt"
    summary_path = os.path.join(summary_folder, summary_file)

    if os.path.exists(summary_path):
        return  # already cached

    st.info(f"❓ لم يُعثر على ملخص لـ `{file_name}`. جارٍ إنشاؤه الآن…")
    original_path = os.path.join(doc_folder, file_name)
    _, content = read_file_content(original_path)

    if not content.strip():
        st.warning(f"⚠️ الملف `{file_name}` فارغ أو غير قابل للقراءة.")
        return

    prompt = f"""
مهمتك هي إنشاء ملخص موجز للمستند التالي. اكتب الملخص كفقرة واحدة متصلة
بدون نقاط أو قوائم. لا تضف مقدمات أو ختامات.

المستند:
{content}
الملخص:
""".strip()

    with st.spinner(f"🚧 توليد ملخص `{file_name}`…"):
        summary = get_deepseek_response_blocking(prompt, api_key)

    if summary:
        try:
            os.makedirs(summary_folder, exist_ok=True)
            full_text = f"ملخص ملف: {file_name}\n\n{summary}"
            with open(summary_path, "w", encoding="utf-8") as f:
                f.write(full_text)
            st.toast(f"✅ تم حفظ ملخص `{file_name}`.", icon="📝")
        except Exception as e:
            st.error(f"❗ فشل حفظ الملخص: {e}")
    else:
        st.error(f"❗ فشل توليد ملخص للملف `{file_name}`.")


def load_all_summaries(summary_folder: str) -> str:
    """Concatenate all summary txt files into a single string."""
    if not os.path.isdir(summary_folder):
        return "لا توجد ملخصات متاحة."

    out = ""
    for fname in sorted(os.listdir(summary_folder)):
        if fname.lower().endswith(".txt"):
            _, txt = read_file_content(os.path.join(summary_folder, fname))
            out += f"---\n{txt}\n---\n\n"
    return out or "لم يُعثر على ملخصات."


def get_document_selection_with_summaries(
    question: str,
    doc_names: List[str],
    rules: str,
    summaries: str,
    api_key: str,
    max_docs: int,
) -> Optional[List[int]]:
    """Ask the LLM to pick the most relevant documents (max *max_docs*)."""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }

    doc_list = "\n".join([f"#{i}: {name}" for i, name in enumerate(doc_names)])

    user_prompt = f"""
مهمتك هي تحديد أنسب الملفات للإجابة على سؤال المستخدم،
بحد أقصى {max_docs} ملف. استخدم القواعد والملخصات.

سؤال المستخدم:
{question}

القواعد الإلزامية:
{rules}

ملخصات الملفات:
{summaries}

قائمة الملفات المتاحة:
{doc_list}

الرد يجب أن يكون أرقام الملفات المختارة بصيغة #N أو #NONE إذا لا شيء مناسب.
""".strip()

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": user_prompt}],
        "stream": False,
    }

    for attempt in range(SELECTION_MAX_RETRIES):
        try:
            resp = requests.post(
                DEEPSEEK_API_URL,
                headers=headers,
                json=payload,
                timeout=45,
            )
            resp.raise_for_status()
            data = resp.json()
            content = data["choices"][0]["message"]["content"]

            found = re.findall(r"#(\d+)", content)

            if found:
                indices = [int(i) for i in found if 0 <= int(i) < len(doc_names)]
                return indices[:max_docs]
            if "#NONE" in content.upper():
                return []
            st.warning(f"⚠️ استجابة غير صالحة من الذكاء الاصطناعي ({attempt+1}/{SELECTION_MAX_RETRIES})")
            time.sleep(1)
        except Exception as e:
            st.error(f"❗ فشل اختيار الوثائق: {e}")
            return None

    st.error("❗ فشل اختيار الوثائق بعد عدة محاولات.")
    return None


def get_deepseek_response_stream(
    prompt: str,
    api_key: str,
    chat_history: List[Dict[str, str]],
    context_doc_count: int,
) -> Generator[str, None, None]:
    """Yield tokens from DeepSeek's streaming endpoint."""
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    messages = chat_history + [{"role": "user", "content": prompt}]
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "stream": True,
    }

    try:
        resp = requests.post(
            DEEPSEEK_API_URL,
            headers=headers,
            json=payload,
            stream=True,
            timeout=60,
        )
        resp.raise_for_status()

        for line in resp.iter_lines():
            if not line:
                continue
            decoded = line.decode("utf-8")
            if not decoded.startswith("data: "):
                continue
            payload = decoded[6:].strip()
            if payload == "[DONE]":
                continue
            try:
                json_data = json.loads(payload)
                delta = json_data.get("choices", [{}])[0].get("delta", {})
                content = delta.get("content", "")
                if content:
                    yield content
            except json.JSONDecodeError:
                continue
    except requests.HTTPError as e:
        err_msg = f"❗ فشل اتصال API (كود {e.response.status_code})."
        if e.response.status_code == 400:
            err_msg += "\nالملف المختار ربما كبير جدًا."
        st.error(err_msg)
        yield err_msg
    except requests.RequestException as e:
        err_msg = f"❗ فشل اتصال الإنترنت: {e}"
        st.error(err_msg)
        yield err_msg
    except Exception as e:
        err_msg = f"❗ خطأ غير متوقع: {e}"
        st.error(err_msg)
        yield err_msg


def build_prompt(
    task_description: str,
    rules: str,
    context: str,
    question: Optional[str] = None,
) -> str:
    """Compose the final LLM prompt."""
    q_section = f"\nسؤال المستخدم:\n---\n{question}\n---\n" if question else ""
    return f"""أنت مساعد قانوني متخصص "RegTech Assistance". اتبع التعليمات بدقة:

قواعد ملزمة:
{rules}

المهمة:
{task_description}

الوثائق:
{context if context else "لم يتم اختيار أو العثور على وثائق."}
{q_section}
"""


# ──────────────────────────────────────
# Streamlit UI
# ──────────────────────────────────────
@safe_run
def main() -> None:
    # --------------------------------------------------------------------- #
    # Page config – must be first Streamlit command
    # --------------------------------------------------------------------- #
    st.set_page_config(
        page_title="RegTech Assistance",
        page_icon="⚖️",
        layout="centered",
    )

    # --------------------------------------------------------------------- #
    # Basic sanity checks
    # --------------------------------------------------------------------- #
    if not LIBRARIES_AVAILABLE:
        st.stop()

    if not API_KEY or not API_KEY.startswith("sk-"):
        st.error("❗ مفتاح DeepSeek غير مضبوط (متغيّر البيئة `DEEPSEEK_API_KEY`).")
        st.stop()

    # --------------------------------------------------------------------- #
    # Session state init
    # --------------------------------------------------------------------- #
    for key, default in {
        "messages": [],
        "doc_names": [],
        "rules_content": "",
        "docs_loaded_first_time": False,
        "action": None,
    }.items():
        if key not in st.session_state:
            st.session_state[key] = default

    # --------------------------------------------------------------------- #
    # Title + separator
    # --------------------------------------------------------------------- #
    st.title("RegTech Assistance")
    st.markdown("<hr/>", unsafe_allow_html=True)

    # --------------------------------------------------------------------- #
    # Main chat container
    # --------------------------------------------------------------------- #
    chat_container = st.container()
    with chat_container:
        # ------------------------------------------------- #
        # 1️⃣ First‑time loading of docs / rules / summaries
        # ------------------------------------------------- #
        if not st.session_state.docs_loaded_first_time:
            with st.spinner("جاري تحميل الوثائق والقواعد..."):
                # Create folders if they are missing (ignore errors on read‑only FS)
                for p in (FOLDER_PATH, SUMMARY_FOLDER_PATH):
                    try:
                        os.makedirs(p, exist_ok=True)
                    except Exception as e:
                        st.warning(f"⚠️ لا يمكن إنشاء المجلد {p}: {e}")

                st.session_state.doc_names = get_document_names_from_folder(FOLDER_PATH)
                st.session_state.rules_content = load_rules_file(RULES_FILE_PATH)
                st.session_state.docs_loaded_first_time = True
                st.toast(
                    f"✅ عُثر على {len(st.session_state.doc_names)} وثيقة.", icon="✅"
                )

            # ----- Auto‑summarize if the flag is on -----
            if USE_AUTO_SUMMARIZATION and st.session_state.doc_names:
                st.write("🧐 فحص الملخصات وإنشاء الناقصة …")
                prog = st.progress(0)
                for i, name in enumerate(st.session_state.doc_names):
                    generate_and_save_summary(
                        name, FOLDER_PATH, SUMMARY_FOLDER_PATH, API_KEY
                    )
                    prog.progress((i + 1) / len(st.session_state.doc_names))
                prog.empty()

        # ------------------------------------------------- #
        # 2️⃣ Action UI – Upload / Summarize
        # ------------------------------------------------- #
        if st.session_state.action == "upload_doc":
            uploaded_file = st.file_uploader(
                "اختر ملفًا لتحميله إلى قاعدة المعرفة",
                type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
            )
            if uploaded_file:
                target_path = os.path.join(FOLDER_PATH, uploaded_file.name)
                try:
                    with open(target_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    st.toast(f"✅ تم رفع `{uploaded_file.name}`", icon="✅")
                    if USE_AUTO_SUMMARIZATION:
                        generate_and_save_summary(
                            uploaded_file.name,
                            FOLDER_PATH,
                            SUMMARY_FOLDER_PATH,
                            API_KEY,
                        )
                except Exception as e:
                    st.error(f"❗ فشل رفع الملف (القرص ربما للقراءة فقط): {e}")

                # Force a reload of the doc list
                st.session_state.docs_loaded_first_time = False
                st.session_state.action = None
                st.rerun()

        if st.session_state.action == "summarize_doc":
            if st.session_state.doc_names:
                selected = st.selectbox(
                    "اختر ملفًا لتلخيصه:",
                    options=st.session_state.doc_names,
                    index=None,
                    placeholder="اختر ملف …",
                )
                if selected:
                    st.session_state.messages.append(
                        {"role": "user", "content": f"لخص لي الملف: `{selected}`"}
                    )
                    st.session_state.action = "processing_summary"
                    st.session_state.summarize_file = selected
                    st.rerun()
            else:
                st.warning("⚠️ لا توجد وثائق لتلخيصها.")
                st.session_state.action = None

        # ------------------------------------------------- #
        # 3️⃣ Display chat history
        # ------------------------------------------------- #
        if not st.session_state.messages:
            st.markdown(
                "<div style='text-align:center;color:#a0a0a0'>مرحباً! كيف يمكنني مساعدتك في وثائقك اليوم؟ ⚖️</div>",
                unsafe_allow_html=True,
            )
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                if msg.get("is_review_notice"):
                    st.info(msg["content"])
                else:
                    st.markdown(msg["content"])

        # ------------------------------------------------- #
        # 4️⃣ Process assistant actions
        # ------------------------------------------------- #
        if st.session_state.action == "processing_summary":
            with st.chat_message("assistant"):
                file_name = st.session_state.summarize_file
                with st.spinner(f"جاري تلخيص `{file_name}`…"):
                    _, text = read_file_content(os.path.join(FOLDER_PATH, file_name))
                    if text:
                        task = "لخص الوثيقة بدقة واذكر النقاط الرئيسية."
                        context = f"المصدر: {file_name}\n\nالمحتوى: {text}"
                        prompt = build_prompt(
                            task, st.session_state.rules_content, context
                        )
                        placeholder = st.empty()
                        full_resp = ""
                        for chunk in get_deepseek_response_stream(
                            prompt, API_KEY, [], 1
                        ):
                            full_resp += chunk
                            placeholder.markdown(full_resp + "▌")
                        placeholder.markdown(full_resp)
                        st.session_state.messages.append(
                            {"role": "assistant", "content": full_resp}
                        )
                    else:
                        err = f"❗ لا يمكن قراءة محتوى `{file_name}`."
                        st.error(err)
                        st.session_state.messages.append(
                            {"role": "assistant", "content": err}
                        )
            st.session_state.action = None
            st.session_state.summarize_file = None
            time.sleep(0.1)
            st.rerun()

        # ------------------------------------------------- #
        # 5️⃣ New user query
        # ------------------------------------------------- #
        if prompt := st.chat_input(
            "اسأل سؤالاً عن مستنداتك …",
            disabled=(st.session_state.action is not None),
        ):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("assistant"):
                final_context = ""
                context_docs: List[Document] = []

                # ---- 2‑step selection (if enabled) ----
                if USE_TWO_STEP_SELECTION and st.session_state.doc_names:
                    with st.spinner("تحليل السؤال واختيار الوثائق…"):
                        all_summaries = load_all_summaries(SUMMARY_FOLDER_PATH)
                        selected_idxs = get_document_selection_with_summaries(
                            prompt,
                            st.session_state.doc_names,
                            st.session_state.rules_content,
                            all_summaries,
                            API_KEY,
                            MAX_DOCS_FOR_SELECTION,
                        )

                    if selected_idxs is None:
                        err = "❗ حدث خطأ أثناء اختيار الوثائق."
                        st.error(err)
                        st.session_state.messages.append(
                            {"role": "assistant", "content": err}
                        )
                        st.rerun()
                    elif not selected_idxs:
                        st.info("⚙️ لم تُحدَّد وثائق ذات صلة – سيُستند الرد إلى القواعد فقط.")
                    else:
                        selected_names = [
                            st.session_state.doc_names[i]
                            for i in selected_idxs
                            if 0 <= i < len(st.session_state.doc_names)
                        ]
                        notice = "سأراجِع الوثائق التالية:\n\n" + "\n".join(
                            f"- `{n}`" for n in selected_names
                        )
                        st.info(notice)
                        st.session_state.messages.append(
                            {
                                "role": "assistant",
                                "content": notice,
                                "is_review_notice": True,
                            }
                        )
                        with st.spinner("تحميل محتوى الوثائق المختارة…"):
                            for n in selected_names:
                                _, txt = read_file_content(os.path.join(FOLDER_PATH, n))
                                if txt:
                                    context_docs.append(
                                        Document(page_content=txt, metadata={"source": n})
                                    )
                        final_context = "\n\n---\n\n".join(
                            f"المصدر: {doc.metadata['source']}\n\nالمحتوى: {doc.page_content}"
                            for doc in context_docs
                        )
                else:
                    st.info("⚙️ اختيار الوثائق معطل – سيتم الاعتماد على القواعد فقط.")
                    final_context = ""

                # ---- Generate the answer ----
                with st.spinner("توليد الرد…"):
                    task = "أجب عن سؤال المستخدم بالاعتماد على الوثائق فقط."
                    full_prompt = build_prompt(
                        task, st.session_state.rules_content, final_context, prompt
                    )
                    placeholder = st.empty()
                    answer = ""
                    history_for_api = [
                        {"role": m["role"], "content": m["content"]}
                        for m in st.session_state.messages[:-1]
                        if not m.get("is_review_notice")
                    ]
                    for chunk in get_deepseek_response_stream(
                        full_prompt, API_KEY, history_for_api, len(context_docs)
                    ):
                        answer += chunk
                        placeholder.markdown(answer + "▌")
                    placeholder.markdown(answer)

                if answer:
                    st.session_state.messages.append(
                        {"role": "assistant", "content": answer}
                    )
                time.sleep(0.1)
                st.rerun()

    # --------------------------------------------------------------------- #
    # Footer – quick actions
    # ---------------------------------------------------------------------- #
    with st.container():
        st.markdown('<div class="footer-actions">', unsafe_allow_html=True)
        st.markdown(
            '<div class="actions-header">الإجراءات السريعة</div>',
            unsafe_allow_html=True,
        )
        col1, col2, col3 = st.columns(3)

        if col1.button("📂 رفع", use_container_width=True):
            st.session_state.action = "upload_doc"
            st.rerun()

        if col2.button("📄 تلخيص", use_container_width=True):
            st.session_state.action = "summarize_doc"
            st.rerun()

        if col3.button("🧹 مسح", use_container_width=True):
            st.session_state.messages = []
            st.session_state.action = None
            st.toast("✅ تم مسح المحادثة", icon="🗑️")
            st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()