from __future__ import annotations # MUST be the first import
import streamlit as st
import os
import requests
from typing import List, Dict, Generator, Optional, Union
import json
import datetime
import re
import time

# --- Library Imports ---
# We import everything here to ensure names are defined.
try:
    from langchain_core.documents import Document
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
    LIBRARIES_AVAILABLE = True
except ImportError as e:
    st.error(f"A required library is not installed. Please check the installation instructions in the original script. Error: {e}")
    LIBRARIES_AVAILABLE = False


# --- Configuration ---
# Set your folder path and DeepSeek API Key here.
# The folder should be in the same directory as this script.
FOLDER_PATH = "./legal_docs"
API_KEY = "sk-cfbf269ccbbc4e94aa69df94c2a25739" # Your DeepSeek API Key
RULES_FILE_NAME = "RULES.txt" # The name of the mandatory rules file.

# --- New Feature Flag ---
# Set to True to enable the new two-step process where the AI first selects
# relevant documents before answering the question.
USE_TWO_STEP_SELECTION = True
SELECTION_MAX_RETRIES = 3 # Number of times to retry the document selection if the AI fails to respond correctly.


# --- Tesseract Installation Check ---
# Pytesseract requires Google's Tesseract-OCR engine. If you get an error,
# follow the installation instructions from the original script and uncomment the line below.
# pytesseract.pytesseract.tesseract_cmd = r'<path_to_your_tesseract_executable>'


# --- DeepSeek API Configuration ---
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

# --- Helper Functions ---

def read_file_content(file: Union[str, st.runtime.uploaded_file_manager.UploadedFile]) -> tuple[str, str]:
    """
    Reads content from a file path or a Streamlit UploadedFile object.
    Supports .pdf, .docx, .txt, and image files.

    Args:
        file: A file path (str) or a Streamlit UploadedFile object.

    Returns:
        A tuple containing the file name and its extracted text content.
    """
    text = ""
    file_name = ""
    try:
        if isinstance(file, str): # If it's a file path
            file_name = os.path.basename(file)
            file_path = file
            file_opener = open(file_path, 'rb')
        else: # If it's a Streamlit UploadedFile
            file_name = file.name
            file_opener = file

        # Handle .pdf files
        if file_name.endswith(".pdf"):
            reader = pypdf.PdfReader(file_opener)
            for page in reader.pages:
                text += page.extract_text() or ""
        # Handle .docx files
        elif file_name.endswith(".docx"):
            doc = docx.Document(file_opener)
            for para in doc.paragraphs:
                text += para.text + "\n"
        # Handle .txt files
        elif file_name.endswith(".txt"):
            # We need to handle the UploadedFile differently for text
            if not isinstance(file, str):
                text = file.getvalue().decode('utf-8')
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    text = f.read()
        # Handle image files (.png, .jpg, .jpeg) using OCR
        elif file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
            try:
                text = pytesseract.image_to_string(Image.open(file_opener))
            except pytesseract.TesseractNotFoundError:
                st.error("Tesseract غير مثبت أو ليس في مسار النظام. يرجى مراجعة التعليمات.")
            except Exception as ocr_e:
                st.error(f"فشل التعرف الضوئي على الحروف لـ {file_name}: {ocr_e}")

        if isinstance(file, str): # Close the file if we opened it
            file_opener.close()

    except Exception as e:
        st.error(f"فشل في قراءة أو معالجة {file_name}: {e}")

    return file_name, text


def load_documents_from_folder(folder_path: str) -> List[Document]:
    """
    Loads all supported files from a folder into a list of Document objects.
    """
    documents = []
    if not os.path.isdir(folder_path):
        st.error(f"المسار المحدد '{folder_path}' ليس دليلاً صالحًا.")
        return []

    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        _, text = read_file_content(file_path)
        if text:
            documents.append(Document(page_content=text, metadata={"source": file_name}))
    return documents


def get_deepseek_response_stream(prompt: str, api_key: str, chat_history: List[Dict[str, str]]) -> Generator[str, None, None]:
    """
    Sends a request to the DeepSeek API and streams the response.
    This is used for the final, user-facing answer.
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    # The prompt is already fully formatted, so we just add it to the history.
    messages = chat_history + [{"role": "user", "content": prompt}]
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "stream": True
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, stream=True)
        response.raise_for_status()
        for chunk in response.iter_lines():
            if chunk:
                decoded_chunk = chunk.decode('utf-8')
                if decoded_chunk.startswith('data: '):
                    try:
                        # Handle potential [DONE] message
                        if decoded_chunk[6:].strip() == '[DONE]':
                            continue
                        json_data = json.loads(decoded_chunk[6:])
                        if 'choices' in json_data and json_data['choices']:
                            content = json_data['choices'][0].get('delta', {}).get('content', '')
                            if content:
                                yield content
                    except json.JSONDecodeError:
                        continue # Ignore malformed chunks
    except requests.exceptions.RequestException as e:
        error_message = f"فشل طلب الواجهة البرمجية: {e}"
        st.error(error_message)
        yield error_message
    except Exception as e:
        error_message = f"حدث خطأ غير متوقع أثناء استدعاء الواجهة البرمجية: {e}"
        st.error(error_message)
        yield error_message


def get_document_selection(question: str, docs: List[Document], api_key: str) -> Optional[List[int]]:
    """
    First step of the two-step process. Asks the AI to select relevant documents.
    This is a non-streaming, hidden call.
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }

    # Create a numbered list of document names for the AI to choose from.
    doc_list_str = "\n".join([f"#{i}: {doc.metadata['source']}" for i, doc in enumerate(docs)])

    selection_prompt = f"""
You are an expert document indexer. Your task is to select up to 3 document indices from the list below that are most relevant to answer the user's question.
You MUST respond ONLY with the indices in the format '#N, #M, #P'. Do not add any other text, explanation, or markdown.
If no documents seem relevant, respond with '#NONE'.

USER'S QUESTION:
---
{question}
---

AVAILABLE DOCUMENTS:
---
{doc_list_str}
---

Your response must be only the selected indices (e.g., #4, #3, #2).
"""

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": selection_prompt}],
        "stream": False # This is a single, non-streamed call
    }

    for attempt in range(SELECTION_MAX_RETRIES):
        try:
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=30)
            response.raise_for_status() # Will raise an exception for 4xx/5xx errors

            response_data = response.json()
            content = response_data['choices'][0]['message']['content']

            # Use regex to robustly find all numbers preceded by a '#'
            indices_found = re.findall(r'#(\d+)', content)

            if indices_found:
                return [int(i) for i in indices_found]
            elif "#NONE" in content:
                return [] # Return an empty list if AI explicitly says none are relevant

            # If response is received but format is wrong, it counts as an attempt
            st.warning(f"استجابة الذكاء الاصطناعي لاختيار المستندات كانت غير منسقة. جاري إعادة المحاولة... (محاولة {attempt + 1}/{SELECTION_MAX_RETRIES})")
            time.sleep(1) # Wait a second before retrying

        except requests.exceptions.RequestException as e:
            # For network/API errors, we fail immediately as requested.
            st.error(f"فشل طلب الواجهة البرمجية لاختيار المستندات: {e}. لا يمكن المتابعة.")
            return None # Indicate a hard failure
        except (KeyError, IndexError, json.JSONDecodeError) as e:
            # For malformed JSON from the server
            st.error(f"فشل في تحليل استجابة اختيار المستندات: {e}. لا يمكن المتابعة.")
            return None # Indicate a hard failure

    st.error(f"فشل في الحصول على اختيار مستند صالح من الذكاء الاصطناعي بعد {SELECTION_MAX_RETRIES} محاولات.")
    return None # Indicate failure after all retries


def build_prompt(task_description: str, rules: str, context: str, question: Optional[str] = None) -> str:
    """Constructs a complete prompt for the LLM based on the task."""
    
    question_section = f"""
    USER'S QUESTION/REQUEST:
    ---
    {question}
    ---
    """ if question else ""

    return f"""
    You are a highly specialized legal assistant, "Mofti Pro+". You must strictly follow the rules provided below in all your responses.

    MANDATORY RULES:
    ---
    {rules}
    ---

    Your task is to perform the following action: {task_description}
    Base your response *only* on the provided context documents. Do not use any external knowledge.
    If the information is not found in the context, state that clearly.
    When answering questions or summarizing, you MUST cite the source file(s) you used from the context.

    CONTEXT DOCUMENTS:
    ---
    {context}
    ---
    {question_section}
    """

# --- Streamlit App UI ---
def main():
    """The main function that runs the Streamlit application."""
    # Set page config
    st.set_page_config(
        page_title="+مفتي بلس",
        page_icon="⚖️",
        initial_sidebar_state="collapsed"
    )
    
    # --- RTL and Custom CSS Injection ---
    st.markdown("""
    <style>
        /* Global RTL setting */
        html, body, [class*="st-"], [data-testid="stAppViewContainer"], [data-testid="stHeader"], [data-testid="stSidebar"] {
            direction: rtl;
        }
        /* Align text content to the right */
        h1, h2, h3, h4, h5, h6, p, li, .stMarkdown, [data-testid="stChatMessageContent"], .stAlert, .stInfo, .stSuccess, .stWarning, .stError, .stToast {
            text-align: right;
        }
        /* Fix for Streamlit's internal layout for chat messages to work in RTL */
        [data-testid="stHorizontalBlock"]:has(> [data-testid="stChatMessageContent"]) {
             flex-direction: row-reverse;
        }
        /* Correct alignment of chat avatars in RTL */
        [data-testid="stChatAvatar"] {
             margin-right: 0;
             margin-left: 1rem;
        }
        /* Make selectbox options appear correctly */
        .stSelectbox div[data-baseweb="select"] {
            text-align: right;
        }
        /* Keep button layout LTR to have icon on the right of text, then reverse it */
        .stButton>button {
            direction: ltr;
            flex-direction: row-reverse;
            justify-content: center;
        }
        /* Style for the smaller action buttons */
        [data-testid="stVerticalBlock"] [data-testid="stHorizontalBlock"] .stButton button {
            white-space: nowrap;
            font-size: 14px;
            padding: 4px 8px;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # --- Main Container for Chat Display ---
    chat_container = st.container()
    with chat_container:
        st.title("⚖️ +مفتي بلس")
        
        # Display a welcome message if the chat is empty
        if "messages" not in st.session_state or not st.session_state.messages:
            st.markdown("مرحباً، كيف يمكنني مساعدتك في مستنداتك القانونية اليوم؟ ⚖️")

        # Display existing chat messages
        if "messages" in st.session_state:
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

    if not LIBRARIES_AVAILABLE:
        st.stop()
    if not API_KEY or "sk-" not in API_KEY:
        st.error("مفتاح واجهة برمجة تطبيقات DeepSeek غير معين بشكل صحيح في أعلى النص البرمجي.")
        st.stop()

    # --- Initialize Session State ---
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "loaded_docs" not in st.session_state:
        st.session_state.loaded_docs = []
    if "docs_loaded_first_time" not in st.session_state:
        st.session_state.docs_loaded_first_time = False
    if "temp_file_content" not in st.session_state:
        st.session_state.temp_file_content = None
    if "selected_docs_map" not in st.session_state:
        st.session_state.selected_docs_map = {}
    # New states for UI control
    if "sidebar_visible" not in st.session_state:
        st.session_state.sidebar_visible = False
    if "show_file_selector" not in st.session_state:
        st.session_state.show_file_selector = False
    if "show_temp_uploader" not in st.session_state:
        st.session_state.show_temp_uploader = False
    if "action" not in st.session_state:
        st.session_state.action = None


    # --- Document Loading (runs only once or when refreshed) ---
    if not st.session_state.docs_loaded_first_time:
        with st.spinner("جاري تحميل المستندات من التخزين الدائم..."):
            if not os.path.exists(FOLDER_PATH):
                os.makedirs(FOLDER_PATH)
            st.session_state.loaded_docs = load_documents_from_folder(FOLDER_PATH)
            
            # Initialize the selection map
            all_doc_names = [doc.metadata['source'] for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME]
            st.session_state.selected_docs_map = {name: True for name in all_doc_names}

            st.session_state.docs_loaded_first_time = True
            if st.session_state.loaded_docs:
                st.toast(f"تم بنجاح تحميل {len(st.session_state.loaded_docs)} مستند.", icon="✅")
            else:
                st.warning("لم يتم العثور على مستندات دائمة. يمكنك تحميلها من لوحة التحكم.")

    # --- Sidebar Toggle Icon (Top Left in RTL) ---
    if st.session_state.sidebar_visible:
        # This button appears at the top of the main page to close the sidebar
        _, col2 = st.columns([20, 1]) 
        with col2:
            if st.button("⬅️", help="إخفاء لوحة التحكم"):
                st.session_state.sidebar_visible = False
                st.rerun()

    # --- Conditional Sidebar (Admin Panel) ---
    if st.session_state.sidebar_visible:
        with st.sidebar:
            st.header("🔐 لوحة التحكم")
            st.info("قم بتحميل الملفات هنا لجعلها متاحة بشكل دائم لجميع الجلسات.")
            # Feature: Permanent File Uploader
            uploaded_files = st.file_uploader(
                "تحميل مستندات دائمة",
                type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
                accept_multiple_files=True,
                help="الملفات التي يتم تحميلها هنا يتم حفظها على الخادم وستكون متاحة في كل مرة تفتح فيها التطبيق."
            )
            if uploaded_files:
                with st.spinner("جاري حفظ ومعالجة الملفات الدائمة..."):
                    for uploaded_file in uploaded_files:
                        with open(os.path.join(FOLDER_PATH, uploaded_file.name), "wb") as f:
                            f.write(uploaded_file.getbuffer())
                st.toast(f"تم تحميل {len(uploaded_files)} ملفات! جاري إعادة التحميل...", icon="🎉")
                st.session_state.docs_loaded_first_time = False
                st.rerun()
            
            if st.button("إخفاء لوحة التحكم"):
                st.session_state.sidebar_visible = False
                st.rerun()

    # --- Handle Actions ---
    if st.session_state.action:
        action = st.session_state.action
        action_placeholder = st.container()

        rules_doc = next((doc for doc in st.session_state.loaded_docs if doc.metadata['source'] == RULES_FILE_NAME), None)
        rules_content = rules_doc.page_content if rules_doc else "No specific rules provided."
        
        prompt_to_send = None
        user_message_content = ""
        context_text = ""
        
        # Update selected_docs based on the map before processing actions
        selectable_docs = [doc for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME]
        selected_docs = [doc for doc in selectable_docs if st.session_state.selected_docs_map.get(doc.metadata['source'])]

        with action_placeholder:
            if action == "summarize_single":
                st.info("اختر ملفًا لتلخيصه.")
                file_to_summarize = st.selectbox("اختر مستندًا:", [d.metadata['source'] for d in selectable_docs])
                if st.button("إنشاء ملخص"):
                    doc_to_summarize = next((d for d in selectable_docs if d.metadata['source'] == file_to_summarize), None)
                    if doc_to_summarize:
                        context_text = f"Source: {doc_to_summarize.metadata['source']}\n\nContent: {doc_to_summarize.page_content}"
                        task = f"Provide a concise summary of the key points, findings, and conclusions from the document '{file_to_summarize}'."
                        prompt_to_send = build_prompt(task, rules_content, context_text)
                        user_message_content = f"الرجاء تلخيص المستند: {file_to_summarize}"

            elif action == "compare_docs":
                st.info("اختر ملفين لمقارنتهما.")
                doc_names = [d.metadata['source'] for d in selectable_docs]
                col1, col2 = st.columns(2)
                file1 = col1.selectbox("المستند الأول:", doc_names, key="comp1")
                file2 = col2.selectbox("المستند الثاني:", doc_names, key="comp2", index=min(1, len(doc_names)-1))
                if st.button("إجراء المقارنة"):
                    if file1 == file2:
                        st.warning("الرجاء اختيار مستندين مختلفين.")
                    else:
                        doc1 = next((d for d in selectable_docs if d.metadata['source'] == file1), None)
                        doc2 = next((d for d in selectable_docs if d.metadata['source'] == file2), None)
                        context_text = f"Document 1 Source: {doc1.metadata['source']}\nContent:\n{doc1.page_content}\n\n---\n\nDocument 2 Source: {doc2.metadata['source']}\nContent:\n{doc2.page_content}"
                        task = f"Compare and contrast the two provided documents ('{file1}' and '{file2}'). Highlight key similarities, differences, and any potential conflicts between them."
                        prompt_to_send = build_prompt(task, rules_content, context_text)
                        user_message_content = f"الرجاء مقارنة '{file1}' و '{file2}'."

            elif action == "define_term":
                st.info("أدخل مصطلحًا قانونيًا لتعريفه بناءً على المستندات المحددة.")
                term = st.text_input("المصطلح المراد تعريفه:")
                if st.button("الحصول على التعريف"):
                    if not term:
                        st.warning("الرجاء إدخال مصطلح.")
                    elif not selected_docs:
                        st.warning("الرجاء تحديد مستند واحد على الأقل من قائمة 🗂️ لتوفير السياق.")
                    else:
                        context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in selected_docs])
                        task = f"Define the legal term '{term}' using only the information available in the provided context documents. If the term is not found, state that."
                        prompt_to_send = build_prompt(task, rules_content, context_text, question=f"Define: {term}")
                        user_message_content = f"عرّف المصطلح: '{term}'"

            elif action == "draft_clause":
                st.info("صف البند القانوني الذي تريد صياغته.")
                description = st.text_area("وصف البند:")
                if st.button("إنشاء مسودة"):
                    if not description:
                        st.warning("الرجاء وصف البند.")
                    elif not selected_docs:
                        st.warning("الرجاء تحديد مستند واحد على الأقل من قائمة 🗂️ لتوفير سياق للأسلوب والمصطلحات.")
                    else:
                        context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in selected_docs])
                        task = f"Draft a legal clause based on the following description: '{description}'. Use the style, terminology, and legal framework found in the provided context documents."
                        prompt_to_send = build_prompt(task, rules_content, context_text, question=f"Draft a clause for: {description}")
                        user_message_content = f"قم بصياغة بند قانوني لي. الوصف: {description}"

        if prompt_to_send and user_message_content:
            st.session_state.messages.append({"role": "user", "content": user_message_content})
            with st.chat_message("assistant"):
                with st.spinner("المساعد القانوني يعمل..."):
                    response_placeholder = st.empty()
                    full_response = ""
                    chat_history_for_api = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1]]
                    for chunk in get_deepseek_response_stream(prompt_to_send, API_KEY, chat_history_for_api):
                        full_response += chunk
                        response_placeholder.markdown(full_response + "▌")
                    response_placeholder.markdown(full_response)
            st.session_state.messages.append({"role": "assistant", "content": full_response})
            st.session_state.action = None # Clear action after completion
            st.rerun()

    # --- UI Panels Toggled by Icons ---
    
    # Panel for Temporary File Uploader
    if st.session_state.show_temp_uploader:
        temp_file = st.file_uploader(
            "تحميل ملف مؤقت لسؤالك التالي",
            type=['pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'],
            help="سيتم استخدام هذا الملف للسؤال التالي فقط ثم يتم تجاهله."
        )
        if temp_file:
            with st.spinner("جاري قراءة الملف المؤقت..."):
                file_name, text = read_file_content(temp_file)
                if text:
                    st.session_state.temp_file_content = Document(page_content=text, metadata={"source": f"ملف مؤقت: {file_name}"})
                    st.info(f"الملف المؤقت جاهز للاستخدام: `{file_name}`. اسأل سؤالك أدناه.")
                else:
                    st.warning(f"تعذرت قراءة المحتوى من `{file_name}`.")
                    st.session_state.temp_file_content = None
            st.session_state.show_temp_uploader = False # Hide uploader after selection
            st.rerun()

    # Panel for Selecting Documents
    if st.session_state.show_file_selector:
        with st.expander("تحديد مستندات للسياق", expanded=True):
            st.caption("اختر الملفات الدائمة التي سيتم تضمينها عند الإجابة على الأسئلة أو تنفيذ الإجراءات.")
            
            selectable_docs = [doc for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME]
            
            if not selectable_docs:
                st.info("لم يتم العثور على مستندات دائمة. قم بتحميل بعضها في لوحة التحكم.")
            else:
                for doc in selectable_docs:
                    doc_name = doc.metadata['source']
                    is_selected = st.checkbox(doc_name, value=st.session_state.selected_docs_map.get(doc_name, True), key=f"cb_{doc_name}")
                    st.session_state.selected_docs_map[doc_name] = is_selected

            if st.button("تم", use_container_width=True):
                st.session_state.show_file_selector = False
                st.rerun()
    
    # --- Chat Input (moved before action bar) ---
    if prompt := st.chat_input("اسأل سؤالاً عن مستنداتك..."):
        # Handle Admin Panel toggle (secretly)
        if prompt.strip() == "1994":
            st.session_state.sidebar_visible = not st.session_state.sidebar_visible
            st.rerun()
        
        # Handle regular chat prompt
        else:
            # Clear any active action/panel when user starts a new chat
            st.session_state.action = None
            st.session_state.show_file_selector = False
            st.session_state.show_temp_uploader = False

            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                spinner_placeholder = st.empty()
                with spinner_placeholder.status("المساعد القانوني يفكر...", expanded=True) as status:
                    
                    rules_doc = next((doc for doc in st.session_state.loaded_docs if doc.metadata['source'] == RULES_FILE_NAME), None)
                    rules_content = rules_doc.page_content if rules_doc else "No specific rules provided. Follow general instructions."
                    
                    context_docs = []
                    # Decide which workflow to use
                    if USE_TWO_STEP_SELECTION:
                        status.update(label="جاري تحليل السؤال للعثور على المستندات ذات الصلة...", state="running")
                        
                        # Get all selectable documents
                        all_selectable_docs = [doc for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME]
                        if st.session_state.get("temp_file_content"):
                            all_selectable_docs.append(st.session_state.temp_file_content)

                        if not all_selectable_docs:
                            st.warning("لا توجد مستندات متاحة للبحث. يرجى تحميل ملفات دائمة أو مؤقتة.")
                            st.stop()

                        # Step 1: Get document selection from AI
                        selected_indices = get_document_selection(prompt, all_selectable_docs, API_KEY)

                        # Handle failure case
                        if selected_indices is None:
                            status.update(label="فشل في اختيار المستندات. يرجى المحاولة مرة أخرى.", state="error", expanded=False)
                            st.session_state.messages.append({"role": "assistant", "content": "لقد واجهت خطأ أثناء محاولة اختيار المستندات ذات الصلة. يرجى طرح سؤالك مرة أخرى."})
                            st.rerun()

                        if not selected_indices:
                            status.update(label="لم يتم العثور على مستندات ذات صلة. سيتم الرد بناءً على القواعد العامة.", state="running")
                        else:
                            # Step 2: Load content of selected documents
                            for i in selected_indices:
                                if 0 <= i < len(all_selectable_docs):
                                    context_docs.append(all_selectable_docs[i])
                                else:
                                    st.warning(f"اختار الذكاء الاصطناعي فهرس مستند غير صالح (#{i})، سيتم تجاهله.")
                            
                            # Display the names of the selected documents being reviewed
                            if context_docs:
                                doc_names = [f"- `{doc.metadata['source']}`" for doc in context_docs]
                                review_message = "أقوم حاليًا بمراجعة المستندات التالية للإجابة على سؤالك:\n\n" + "\n".join(doc_names)
                                status.update(label=review_message, state="running")
                            else:
                                status.update(label="تحضير الإجابة...", state="running")
                    else:
                        # Original workflow: use pre-selected documents
                        status.update(label="استخدام المستندات المحددة مسبقًا للسياق...", state="running")
                        context_docs = [doc for doc in st.session_state.loaded_docs if doc.metadata['source'] != RULES_FILE_NAME and st.session_state.selected_docs_map.get(doc.metadata['source'])]
                        if st.session_state.get("temp_file_content"):
                            context_docs.append(st.session_state.temp_file_content)

                    if not context_docs and USE_TWO_STEP_SELECTION:
                        # This can happen if selection returns [] or if all indices were invalid.
                        # We still proceed, but the AI will be told the context is empty.
                        st.write("لا توجد مستندات سياق للاستخدام. سيذكر المساعد أنه لا يمكنه العثور على المعلومات.")

                    # Build the final prompt for the answer
                    context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in context_docs])
                    task = "Answer the user's question based on the provided context."
                    prompt_to_send = build_prompt(task, rules_content, context_text, prompt)

                    status.update(label="جاري إنشاء الرد...", state="running")
                
                # This part is now outside the status UI
                spinner_placeholder.empty() # Remove the status UI
                response_placeholder = st.empty()
                full_response = ""
                chat_history_for_api = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1]]

                for chunk in get_deepseek_response_stream(prompt_to_send, API_KEY, chat_history_for_api):
                    full_response += chunk
                    response_placeholder.markdown(full_response + "▌")
                response_placeholder.markdown(full_response)

            st.session_state.messages.append({"role": "assistant", "content": full_response})
            # Clear the temporary file after it has been used
            st.session_state.temp_file_content = None
            st.rerun()

    # --- Container for Action Bar (at the end of the script's flow) ---
    with st.container():
        st.markdown("---")
        
        cols = st.columns(7)
        
        # Dictionary mapping icons to captions and action keys
        button_actions = {
            "📎": ("إرفاق", "show_temp_uploader"),
            "🗂️": ("تحديد", "show_file_selector"),
            "📄": ("ملخص", "summarize_single"),
            "🔄": ("مقارنة", "compare_docs"),
            "📖": ("تعريف", "define_term"),
            "✍️": ("صياغة", "draft_clause"),
            "🧹": ("مسح", "clear_chat")
        }

        # Create buttons with combined icon and text
        for i, (icon, (caption, action_key)) in enumerate(button_actions.items()):
            with cols[i]:
                button_label = f"{icon} {caption}"
                if st.button(button_label, help=f"إجراء {caption}", use_container_width=True):
                    if action_key == "clear_chat":
                        st.session_state.messages = []
                        st.session_state.action = None
                        st.session_state.show_file_selector = False
                        st.session_state.show_temp_uploader = False
                        st.toast("تم مسح سجل المحادثة!", icon="🧹")
                        st.rerun()
                    elif action_key == "show_file_selector":
                        st.session_state.show_file_selector = not st.session_state.show_file_selector
                        st.session_state.show_temp_uploader = False
                        st.session_state.action = None
                    elif action_key == "show_temp_uploader":
                        st.session_state.show_temp_uploader = not st.session_state.show_temp_uploader
                        st.session_state.show_file_selector = False
                        st.session_state.action = None
                    else:
                        st.session_state.action = action_key
                        st.session_state.show_file_selector = False
                        st.session_state.show_temp_uploader = False
                    st.rerun()

if __name__ == "__main__":
    main()