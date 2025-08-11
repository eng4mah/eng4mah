from __future__ import annotations # MUST be the first import
import streamlit as st
import os
import requests
from typing import List, Dict, Generator, Optional, Union, Tuple
import json
import datetime
import re
import time

# --- Library Imports ---
try:
    from langchain_core.documents import Document
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
    LIBRARIES_AVAILABLE = True
except ImportError as e:
    st.error(f"مكتبة ناقصة: {e}")
    LIBRARIES_AVAILABLE = False

# --- Configuration ---
FOLDER_PATH = "./legal_docs"
API_KEY = "sk-cfbf269ccbbc4e94aa69df94c2a25739"
RULES_FILE_NAME = "RULES.txt"
USE_TWO_STEP_SELECTION = True
USE_LAZY_LOADING = True
SELECTION_MAX_RETRIES = 3
MAX_DOCS_FOR_SELECTION = 2

DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

# --- CSS for 3-part layout, Arabic, red lines, sticky areas ---
st.markdown("""
<style>
html, body, [class*="st-"] {
    direction: rtl !important;
    text-align: right !important;
}
.st-emotion-cache-uf99v8 {
    direction: rtl !important;
    text-align: right !important;
}
.stChatInputContainer {
    position: sticky;
    bottom: 0;
    z-index: 100;
    background: #fff;
    border-top: 2px solid #E53935;
}
#quick-actions {
    position: sticky;
    bottom: 0;
    z-index: 101;
    background: #1a1a1a;
    border-top: 2px solid #E53935;
    padding-bottom: 2rem;
}
hr {
    border: 2px solid #E53935;
    margin: 0.6em 0;
}
.actions-header {
    text-align: center;
    font-weight: bold;
    color: #E53935;
    font-size: 1.1em;
    margin-top: 0.5em;
    margin-bottom: 1em;
    letter-spacing: 1px;
}
h1 { text-align: center; }
</style>
""", unsafe_allow_html=True)

# --- Helper Functions (كما هي من كودك السابق) ---
def read_file_content(file_path: str) -> Tuple[str, str]:
    text = ""
    file_name = os.path.basename(file_path)
    try:
        with open(file_path, 'rb') as f:
            if file_name.endswith(".pdf"):
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            elif file_name.endswith(".docx"):
                doc = docx.Document(f)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif file_name.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as text_f:
                    text = text_f.read()
            elif file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
                try:
                    text = pytesseract.image_to_string(Image.open(f))
                except pytesseract.TesseractNotFoundError:
                    st.error("مطلوب تثبيت Tesseract.")
                except Exception as ocr_e:
                    st.error(f"فشل قراءة الصورة {file_name}: {ocr_e}")
    except Exception as e:
        st.error(f"خطأ قراءة الملف {file_name}: {e}")
    return file_name, text

def get_document_names_from_folder(folder_path: str) -> List[str]:
    if not os.path.isdir(folder_path):
        return []
    supported_files = []
    for file_name in os.listdir(folder_path):
        if file_name != RULES_FILE_NAME and file_name.lower().endswith(('.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg')):
            supported_files.append(file_name)
    return supported_files

def load_rules_file(folder_path: str, rules_filename: str) -> str:
    rules_path = os.path.join(folder_path, rules_filename)
    if os.path.exists(rules_path):
        _, content = read_file_content(rules_path)
        return content
    return "لا توجد تعليمات خاصة. اتبع التعليمات العامة."

def get_deepseek_response_stream(prompt: str, api_key: str, chat_history: List[Dict[str, str]], context_doc_count: int) -> Generator[str, None, None]:
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    messages = chat_history + [{"role": "user", "content": prompt}]
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "stream": True
    }
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, stream=True, timeout=60)
        response.raise_for_status()
        for chunk in response.iter_lines():
            if chunk:
                decoded_chunk = chunk.decode('utf-8')
                if decoded_chunk.startswith('data: '):
                    try:
                        if decoded_chunk[6:].strip() == '[DONE]':
                            continue
                        json_data = json.loads(decoded_chunk[6:])
                        if 'choices' in json_data and json_data['choices']:
                            content = json_data['choices'][0].get('delta', {}).get('content', '')
                            if content:
                                yield content
                    except json.JSONDecodeError:
                        continue
    except requests.exceptions.HTTPError as e:
        error_message = f"**فشل الاتصال بالـ API برمز {e.response.status_code}.**"
        if e.response.status_code == 400:
            error_message += f"\n\nغالباً النص طويل جداً أو عدد الوثائق كثير (**{context_doc_count}**). جرب تقليلها."
        else:
            error_message += f"\n\n**السبب:** {e.response.reason}\n**التفاصيل:**\n```\n{e.response.text}\n```"
        st.error(error_message)
        yield error_message
    except requests.exceptions.RequestException as e:
        error_message = f"فشل الاتصال بالإنترنت: {e}"
        st.error(error_message)
        yield error_message
    except Exception as e:
        error_message = f"حدث خطأ غير متوقع: {e}"
        st.error(error_message)
        yield error_message

def get_document_selection(question: str, doc_names: List[str], api_key: str) -> Optional[List[int]]:
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    doc_list_str = "\n".join([f"#{i}: {name}" for i, name in enumerate(doc_names)])
    selection_prompt = f"""
حدد حتى {MAX_DOCS_FOR_SELECTION} ملف من القائمة التالية الأكثر صلة بسؤال المستخدم.
يجب أن ترد فقط بأرقام الملفات بصيغة: '#N, #M'. إذا لا يوجد، اكتب '#NONE'.

سؤال المستخدم:
---
{question}
---
الملفات المتوفرة:
---
{doc_list_str}
---
"""
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": selection_prompt}],
        "stream": False
    }
    for attempt in range(SELECTION_MAX_RETRIES):
        try:
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=30)
            if response.status_code == 400:
                st.error(f"خطأ في طلب API (400). لا يمكن إعادة المحاولة. التفاصيل: {response.text}")
                return None
            response.raise_for_status()
            response_data = response.json()
            content = response_data['choices'][0]['message']['content']
            indices_found = re.findall(r'#(\d+)', content)
            if indices_found:
                return [int(i) for i in indices_found]
            elif "#NONE" in content.upper():
                return []
            st.warning(f"رد الذكاء الاصطناعي غير صحيح. إعادة المحاولة... ({attempt+1}/{SELECTION_MAX_RETRIES})")
            time.sleep(1)
        except Exception:
            return None
    st.error("فشل اختيار الوثائق من الذكاء الاصطناعي بعد عدة محاولات.")
    return None

def build_prompt(task_description: str, rules: str, context: str, question: Optional[str] = None) -> str:
    question_section = f"\nسؤال المستخدم:\n---\n{question}\n---\n" if question else ""
    return f"""
أنت مساعد قانوني متخصص "مفتي بلس". يجب عليك اتباع التعليمات التالية بدقة:

قواعد ملزمة:
---
{rules}
---

المهمة: {task_description}
اعتمد فقط على الوثائق المقدمة. لا تستخدم معلومات خارجية.
إذا لم تجد المعلومة، صرّح بذلك.
يجب ذكر مصدر كل معلومة.

الوثائق:
---
{context if context else "لم يتم اختيار أو العثور على وثائق."}
---
{question_section}
"""

# --- Streamlit 3-part layout ---
def main():
    st.set_page_config(
        page_title="+ مفتي بلس",
        page_icon="⚖️",
        layout="centered"
    )
    if not LIBRARIES_AVAILABLE:
        st.stop()
    if not API_KEY or "sk-" not in API_KEY:
        st.error("لم يتم ضبط مفتاح DeepSeek API بشكل صحيح.")
        st.stop()

    # --- Session state ---
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "doc_names" not in st.session_state:
        st.session_state.doc_names = []
    if "rules_content" not in st.session_state:
        st.session_state.rules_content = ""
    if "docs_loaded_first_time" not in st.session_state:
        st.session_state.docs_loaded_first_time = False

    # --- 1. Title Part ---
    with st.container():
        st.title("+ مفتي بلس")
    st.markdown("<hr />", unsafe_allow_html=True)

    # --- 2. Chat Part ---
    with st.container():
        if USE_LAZY_LOADING and not st.session_state.docs_loaded_first_time:
            with st.spinner("جاري تحميل قائمة الوثائق والقواعد..."):
                if not os.path.exists(FOLDER_PATH):
                    os.makedirs(FOLDER_PATH)
                st.session_state.doc_names = get_document_names_from_folder(FOLDER_PATH)
                st.session_state.rules_content = load_rules_file(FOLDER_PATH, RULES_FILE_NAME)
                st.session_state.docs_loaded_first_time = True
                st.toast(f"تم العثور على {len(st.session_state.doc_names)} وثيقة.", icon="✅")

        # Chat display
        if not st.session_state.messages:
            st.markdown("<div style='text-align: center; color: #a0a0a0;'>مرحباً، كيف أساعدك في مستنداتك القانونية اليوم؟ ⚖️</div>", unsafe_allow_html=True)
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                if message.get("is_review_notice"):
                    st.info(message["content"])
                else:
                    st.markdown(message["content"])

        # Chat input always at the bottom of this section
        prompt = st.chat_input("اسأل سؤالاً عن مستنداتك...")
        if prompt:
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("assistant"):
                context_docs_content = []
                final_context_text = ""
                if USE_TWO_STEP_SELECTION and st.session_state.doc_names:
                    with st.spinner("جاري تحليل سؤالك وتحديد الوثائق المناسبة..."):
                        selected_indices = get_document_selection(prompt, st.session_state.doc_names, API_KEY)
                    if selected_indices is None:
                        st.error("حدث خطأ أثناء اختيار الوثائق. أعد المحاولة.")
                        st.session_state.messages.append({"role": "assistant", "content": "حدث خطأ أثناء اختيار الوثائق."})
                        st.rerun()
                    elif not selected_indices:
                        st.info("لا توجد وثائق ذات صلة واضحة. سأجيب بناءً على القواعد فقط.")
                        final_context_text = ""
                    else:
                        selected_names = [st.session_state.doc_names[i] for i in selected_indices if 0 <= i < len(st.session_state.doc_names)]
                        review_notice = "سأراجع الملفات التالية:\n\n" + "\n".join([f"- `{name}`" for name in selected_names])
                        st.info(review_notice)
                        st.session_state.messages.append({"role": "assistant", "content": review_notice, "is_review_notice": True})
                        with st.spinner("جاري تحميل محتوى الملفات المختارة..."):
                            for name in selected_names:
                                file_path = os.path.join(FOLDER_PATH, name)
                                _, text = read_file_content(file_path)
                                if text:
                                    context_docs_content.append(Document(page_content=text, metadata={"source": name}))
                        final_context_text = "\n\n---\n\n".join([f"المصدر: {doc.metadata['source']}\n\nالمحتوى: {doc.page_content}" for doc in context_docs_content])
                else:
                    with st.spinner("جاري تحميل كل الوثائق المتاحة..."):
                        for name in st.session_state.doc_names:
                            file_path = os.path.join(FOLDER_PATH, name)
                            _, text = read_file_content(file_path)
                            if text:
                                context_docs_content.append(Document(page_content=text, metadata={"source": name}))
                        final_context_text = "\n\n---\n\n".join([f"المصدر: {doc.metadata['source']}\n\nالمحتوى: {doc.page_content}" for doc in context_docs_content])
                # Step 2: Final Answer
                with st.spinner("جاري توليد الرد..."):
                    task = "أجب عن سؤال المستخدم بناءً على الوثائق فقط."
                    final_prompt = build_prompt(task, st.session_state.rules_content, final_context_text, prompt)
                    response_placeholder = st.empty()
                    full_response = ""
                    chat_history_for_api = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1] if not m.get("is_review_notice")]
                    stream_generator = get_deepseek_response_stream(final_prompt, API_KEY, chat_history_for_api, len(context_docs_content))
                    for chunk in stream_generator:
                        full_response += chunk
                        response_placeholder.markdown(full_response + "▌")
                    response_placeholder.markdown(full_response)
                if full_response:
                    st.session_state.messages.append({"role": "assistant", "content": full_response})
                time.sleep(0.15)
                st.rerun()

    st.markdown("<hr />", unsafe_allow_html=True)

    # --- 3. Quick Actions Part (fixed) ---
    with st.container():
        st.markdown('<div id="quick-actions">', unsafe_allow_html=True)
        st.markdown("<div class='actions-header'>الإجراءات السريعة</div>", unsafe_allow_html=True)
        action_cols = st.columns(5)
        action_buttons = {
            "📄 ملخص": "summarize_single",
            "🔄 مقارنة": "compare_docs",
            "📖 تعريف": "define_term",
            "✍️ صياغة": "draft_clause",
            "🧹 مسح": "clear_chat"
        }
        for i, (label, action) in enumerate(action_buttons.items()):
            if action_cols[i].button(label, use_container_width=True):
                if action == "clear_chat":
                    st.session_state.messages = []
                    st.toast("تم مسح المحادثة 🧹")
                    st.rerun()
                else:
                    st.info(f"تم اختيار الإجراء: '{label}'، يرجى إدخال التفاصيل في الأسفل.")
        st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
