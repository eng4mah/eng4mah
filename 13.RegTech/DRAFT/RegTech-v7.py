from __future__ import annotations # MUST be the first import
import streamlit as st
import os
import requests
from typing import List, Dict, Generator, Optional, Union, Tuple
import json
import datetime
import re
import time

# --- Library Imports ---
# We import everything here to ensure names are defined.
try:
    from langchain_core.documents import Document
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
    LIBRARIES_AVAILABLE = True
except ImportError as e:
    st.error(f"A required library is not installed. Please check the installation instructions in the original script. Error: {e}")
    LIBRARIES_AVAILABLE = False


# --- Configuration ---
# Set your folder path and DeepSeek API Key here.
# The folder should be in the same directory as this script.
FOLDER_PATH = "./legal_docs"
API_KEY = "sk-cfbf269ccbbc4e94aa69df94c2a25739" # Your DeepSeek API Key
RULES_FILE_NAME = "RULES.txt" # The name of the mandatory rules file.

# --- New Feature Flags ---
# Set to True to enable the new two-step process where the AI first selects
# relevant documents before answering the question.
USE_TWO_STEP_SELECTION = True
# Lazy load documents: only load file names on startup, and file content when selected by the AI.
USE_LAZY_LOADING = True
SELECTION_MAX_RETRIES = 3 # Number of times to retry the document selection if the AI fails to respond correctly.
MAX_DOCS_FOR_SELECTION = 2 # The maximum number of documents the AI is allowed to select.


# --- Tesseract Installation Check ---
# Pytesseract requires Google's Tesseract-OCR engine. If you get an error,
# follow the installation instructions from the original script and uncomment the line below.
# pytesseract.pytesseract.tesseract_cmd = r'<path_to_your_tesseract_executable>'


# --- DeepSeek API Configuration ---
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

# --- Helper Functions ---

def read_file_content(file_path: str) -> Tuple[str, str]:
    """
    Reads content from a file path.
    Supports .pdf, .docx, .txt, and image files.

    Args:
        file_path: The full path to the file.

    Returns:
        A tuple containing the file name and its extracted text content.
    """
    text = ""
    file_name = os.path.basename(file_path)
    try:
        with open(file_path, 'rb') as f:
            if file_name.endswith(".pdf"):
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            elif file_name.endswith(".docx"):
                doc = docx.Document(f)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif file_name.endswith(".txt"):
                # We need to re-open in text mode after reading bytes
                with open(file_path, 'r', encoding='utf-8') as text_f:
                    text = text_f.read()
            elif file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
                try:
                    text = pytesseract.image_to_string(Image.open(f))
                except pytesseract.TesseractNotFoundError:
                    st.error("Tesseract is not installed or not in your system's PATH. Please check the instructions.")
                except Exception as ocr_e:
                    st.error(f"OCR failed for {file_name}: {ocr_e}")
    except Exception as e:
        st.error(f"Failed to read or process {file_name}: {e}")

    return file_name, text

def get_document_names_from_folder(folder_path: str) -> List[str]:
    """
    Gets a list of all supported file names from a folder, excluding the RULES file.
    """
    if not os.path.isdir(folder_path):
        return []
    
    supported_files = []
    for file_name in os.listdir(folder_path):
        if file_name != RULES_FILE_NAME and file_name.lower().endswith(('.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg')):
            supported_files.append(file_name)
    return supported_files

def load_rules_file(folder_path: str, rules_filename: str) -> str:
    """
    Loads the content of the rules file.
    """
    rules_path = os.path.join(folder_path, rules_filename)
    if os.path.exists(rules_path):
        _, content = read_file_content(rules_path)
        return content
    return "No specific rules provided. Follow general instructions."


def get_deepseek_response_stream(prompt: str, api_key: str, chat_history: List[Dict[str, str]], context_doc_count: int) -> Generator[str, None, None]:
    """
    Sends a request to the DeepSeek API and streams the response.
    This is used for the final, user-facing answer.
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    messages = chat_history + [{"role": "user", "content": prompt}]
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "stream": True
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, stream=True, timeout=60)
        response.raise_for_status() # Will raise HTTPError for 4xx/5xx responses
        for chunk in response.iter_lines():
            if chunk:
                decoded_chunk = chunk.decode('utf-8')
                if decoded_chunk.startswith('data: '):
                    try:
                        if decoded_chunk[6:].strip() == '[DONE]':
                            continue
                        json_data = json.loads(decoded_chunk[6:])
                        if 'choices' in json_data and json_data['choices']:
                            content = json_data['choices'][0].get('delta', {}).get('content', '')
                            if content:
                                yield content
                    except json.JSONDecodeError:
                        continue # Ignore malformed chunks
    except requests.exceptions.HTTPError as e:
        # Handle specific HTTP errors, like "Bad Request"
        error_message = f"**API Request Failed with status {e.response.status_code}.**"
        if e.response.status_code == 400:
            error_message += f"\n\nThis is often a 'Bad Request' error, which can be caused by sending too much text (e.g., too many documents). You requested content from **{context_doc_count}** document(s). Try selecting fewer documents or asking a more specific question."
        else:
            error_message += f"\n\n**Reason:** {e.response.reason}\n**Details:**\n```\n{e.response.text}\n```"
        st.error(error_message)
        yield error_message
    except requests.exceptions.RequestException as e:
        error_message = f"API Request Failed: {e}"
        st.error(error_message)
        yield error_message
    except Exception as e:
        error_message = f"An unexpected error occurred during the API call: {e}"
        st.error(error_message)
        yield error_message


def get_document_selection(question: str, doc_names: List[str], api_key: str) -> Optional[List[int]]:
    """
    First step of the two-step process. Asks the AI to select relevant documents by name.
    This is a non-streaming, hidden call.
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }

    doc_list_str = "\n".join([f"#{i}: {name}" for i, name in enumerate(doc_names)])

    selection_prompt = f"""
You are an expert document indexer. Your task is to select up to {MAX_DOCS_FOR_SELECTION} document indices from the list below that are most relevant to answer the user's question.
You MUST respond ONLY with the indices in the format '#N, #M, #P'. Do not add any other text, explanation, or markdown.
If no documents seem relevant, respond with '#NONE'.

USER'S QUESTION:
---
{question}
---

AVAILABLE DOCUMENTS:
---
{doc_list_str}
---

Your response must be only the selected indices (e.g., #4, #3, #2).
"""

    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": selection_prompt}],
        "stream": False
    }

    for attempt in range(SELECTION_MAX_RETRIES):
        try:
            response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, timeout=30)
            
            # Immediately check for non-retriable HTTP errors
            if response.status_code == 400:
                st.error(f"API Bad Request (400) during document selection. This is not retriable. Details: {response.text}")
                return None
            response.raise_for_status()

            response_data = response.json()
            content = response_data['choices'][0]['message']['content']

            indices_found = re.findall(r'#(\d+)', content)

            if indices_found:
                return [int(i) for i in indices_found]
            elif "#NONE" in content.upper():
                return [] 

            st.warning(f"AI response for document selection was not formatted correctly. Retrying... (Attempt {attempt + 1}/{SELECTION_MAX_RETRIES})")
            time.sleep(1)

        except requests.exceptions.HTTPError as e:
            st.error(f"HTTP Error during document selection: {e}. This is not retriable.")
            return None 
        except requests.exceptions.RequestException as e:
            st.error(f"Network error during document selection: {e}. This is not retriable.")
            return None
        except (KeyError, IndexError, json.JSONDecodeError) as e:
            st.error(f"Failed to parse the document selection response: {e}. This is not retriable.")
            return None

    st.error(f"Failed to get a valid document selection from the AI after {SELECTION_MAX_RETRIES} attempts. The AI failed to follow the requested format.")
    return None


def build_prompt(task_description: str, rules: str, context: str, question: Optional[str] = None) -> str:
    """Constructs a complete prompt for the LLM based on the task."""
    
    question_section = f"""
USER'S QUESTION/REQUEST:
---
{question}
---
""" if question else ""

    return f"""
You are a highly specialized legal assistant, "Mofti Pro+". You must strictly follow the rules provided below in all your responses.

MANDATORY RULES:
---
{rules}
---

Your task is to perform the following action: {task_description}
Base your response *only* on the provided context documents. Do not use any external knowledge.
If the information is not found in the context, state that clearly.
When answering questions or summarizing, you MUST cite the source file(s) you used from the context.

CONTEXT DOCUMENTS:
---
{context if context else "No context documents were provided or selected for this query."}
---
{question_section}
"""

# --- Streamlit App UI ---
def main():
    """The main function that runs the Streamlit application."""
    st.set_page_config(
        page_title="+ مفتي بلس",
        page_icon="⚖️",
        layout="centered"
    )
    
    # --- RTL and Custom CSS Injection ---
    st.markdown("""
    <style>
        /* Global RTL setting */
        html, body, [class*="st-"], [data-testid="stAppViewContainer"], [data-testid="stHeader"], [data-testid="stSidebar"] {
            direction: rtl;
        }
        /* Align text content to the right */
        h1, h2, h3, h4, h5, h6, p, li, .stMarkdown, [data-testid="stChatMessageContent"], .stAlert, .stInfo, .stSuccess, .stWarning, .stError, .stToast {
            text-align: right;
        }
        /* Fix for Streamlit's internal layout for chat messages to work in RTL */
        [data-testid="stHorizontalBlock"]:has(> [data-testid="stChatMessageContent"]) {
             flex-direction: row-reverse;
        }
        /* Correct alignment of chat avatars in RTL */
        [data-testid="stChatAvatar"] {
             margin-right: 0;
             margin-left: 1rem;
        }
        /* Keep button layout LTR to have icon on the right of text, then reverse it */
        .stButton>button {
            direction: ltr;
            flex-direction: row-reverse;
            justify-content: center;
        }
        /* Main container with red glow effect */
        .chat-action-container {
            border: 1px solid #4a0000;
            border-radius: 15px;
            padding: 1rem;
            box-shadow: 0 0 15px rgba(255, 0, 0, 0.5);
            background-color: #1a1a1a; /* Dark background for contrast */
        }
        /* Add some space below the title */
        h1 {
            margin-bottom: 1.5rem;
        }
        /* Style for the "Quick Actions" header */
        .actions-header {
            text-align: center;
            font-weight: bold;
            color: #cccccc;
            margin-top: 1rem;
            margin-bottom: 1rem;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("+مفتي بلس")

    if not LIBRARIES_AVAILABLE:
        st.stop()
    if not API_KEY or "sk-" not in API_KEY:
        st.error("DeepSeek API Key is not correctly set at the top of the script.")
        st.stop()

    # --- Initialize Session State ---
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "doc_names" not in st.session_state:
        st.session_state.doc_names = []
    if "rules_content" not in st.session_state:
        st.session_state.rules_content = ""
    if "docs_loaded_first_time" not in st.session_state:
        st.session_state.docs_loaded_first_time = False

    # --- Document & Rules Loading (Lazy Loading) ---
    if USE_LAZY_LOADING and not st.session_state.docs_loaded_first_time:
        with st.spinner("Loading document index and rules..."):
            if not os.path.exists(FOLDER_PATH):
                os.makedirs(FOLDER_PATH)
            st.session_state.doc_names = get_document_names_from_folder(FOLDER_PATH)
            st.session_state.rules_content = load_rules_file(FOLDER_PATH, RULES_FILE_NAME)
            st.session_state.docs_loaded_first_time = True
            st.toast(f"Found {len(st.session_state.doc_names)} documents. Ready to assist.", icon="✅")

    # --- Main Container for Chat and Actions ---
    st.markdown('<div class="chat-action-container">', unsafe_allow_html=True)

    # --- Chat Display Area ---
    chat_display_area = st.container()
    with chat_display_area:
        if not st.session_state.messages:
            st.markdown("<div style='text-align: center; color: #a0a0a0; margin-bottom: 1rem;'>مرحباً، كيف يمكنني مساعدتك في مستنداتك القانونية اليوم؟ ⚖️</div>", unsafe_allow_html=True)
        
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                # If it's the special "reviewing files" message, display it differently
                if message.get("is_review_notice"):
                    st.info(message["content"])
                else:
                    st.markdown(message["content"])

    # --- Quick Actions Section ---
    st.markdown("<div class='actions-header'>الاجراءات السريعة</div>", unsafe_allow_html=True)
    action_cols = st.columns(5)
    action_buttons = {
        "📄 ملخص": "summarize_single",
        "🔄 مقارنة": "compare_docs",
        "📖 تعريف": "define_term",
        "✍️ صياغة": "draft_clause",
        "🧹 مسح": "clear_chat"
    }

    for i, (label, action) in enumerate(action_buttons.items()):
        if action_cols[i].button(label, use_container_width=True):
            if action == "clear_chat":
                st.session_state.messages = []
                st.toast("Chat history cleared!", icon="🧹")
                st.rerun()
            else:
                # This part would need a more complex implementation to handle actions
                # without a direct chat prompt, which is outside the current refactor scope.
                # For now, we can just acknowledge the action.
                st.info(f"Action '{label}' selected. Please provide details in the chat box below.")

    st.markdown('</div>', unsafe_allow_html=True) # Close the main container div

    # --- Chat Input Logic ---
    if prompt := st.chat_input("اسأل سؤالاً عن مستنداتك..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("assistant"):
            # This is the container for the entire response generation flow
            response_flow_container = st.container()
            
            with response_flow_container:
                context_docs_content = []
                final_context_text = ""
                
                # --- Step 1: Select Relevant Documents (if enabled) ---
                if USE_TWO_STEP_SELECTION and st.session_state.doc_names:
                    with st.spinner("Analyzing question to find relevant documents..."):
                        selected_indices = get_document_selection(prompt, st.session_state.doc_names, API_KEY)

                    if selected_indices is None:
                        # Hard failure in selection
                        st.error("Could not proceed with document selection due to an API or parsing error.")
                        st.session_state.messages.append({"role": "assistant", "content": "I encountered an error trying to select relevant documents. Please try your question again."})
                        st.rerun()
                    
                    elif not selected_indices:
                        st.info("No specific documents seemed relevant to your question. I will answer based on general knowledge and rules.")
                        final_context_text = "" # No context to add
                    
                    else:
                        # --- Step 1.5: Load selected documents and post notice ---
                        selected_names = [st.session_state.doc_names[i] for i in selected_indices if 0 <= i < len(st.session_state.doc_names)]
                        
                        # Post the "reviewing files" notice
                        review_notice = "سأقوم بمراجعة الملفات التالية:\n\n" + "\n".join([f"- `{name}`" for name in selected_names])
                        st.info(review_notice)
                        # Add to history as a special, non-LLM message
                        st.session_state.messages.append({"role": "assistant", "content": review_notice, "is_review_notice": True})
                        
                        with st.spinner("Loading content of selected documents..."):
                            for name in selected_names:
                                file_path = os.path.join(FOLDER_PATH, name)
                                _, text = read_file_content(file_path)
                                if text:
                                    context_docs_content.append(Document(page_content=text, metadata={"source": name}))
                        
                        final_context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in context_docs_content])

                else:
                    # Fallback or if selection is disabled: use all documents
                    with st.spinner("Loading all available documents for context..."):
                        for name in st.session_state.doc_names:
                            file_path = os.path.join(FOLDER_PATH, name)
                            _, text = read_file_content(file_path)
                            if text:
                                context_docs_content.append(Document(page_content=text, metadata={"source": name}))
                        final_context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in context_docs_content])

                # --- Step 2: Generate the Final Answer ---
                with st.spinner("Generating response..."):
                    task = "Answer the user's question based on the provided context."
                    final_prompt = build_prompt(task, st.session_state.rules_content, final_context_text, prompt)
                    
                    response_placeholder = st.empty()
                    full_response = ""
                    chat_history_for_api = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages[:-1] if not m.get("is_review_notice")]
                    
                    # Get the final answer stream
                    stream_generator = get_deepseek_response_stream(final_prompt, API_KEY, chat_history_for_api, len(context_docs_content))
                    
                    for chunk in stream_generator:
                        full_response += chunk
                        response_placeholder.markdown(full_response + "▌")
                    response_placeholder.markdown(full_response)

                # Append the final, complete response to history and rerun
                if full_response:
                    st.session_state.messages.append({"role": "assistant", "content": full_response})
                
                # Short delay to ensure user sees the final message before rerun
                time.sleep(0.1)
                st.rerun()

if __name__ == "__main__":
    main()
