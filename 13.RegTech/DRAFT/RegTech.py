# main_app.py
from __future__ import annotations # MUST be the first import
import streamlit as st
import os
import requests
from typing import List, Dict, Generator
import json

# --- Configuration ---
# Set your folder path and DeepSeek API Key here.
# The folder should be in the same directory as this script.
FOLDER_PATH = "./legal_docs"
API_KEY = "sk-cfbf269ccbbc4e94aa69df94c2a25739" # Your DeepSeek API Key

# --- Installation ---
# Before running, install all required libraries. Open your terminal and run:
# pip install streamlit requests pypdf python-docx Pillow pytesseract

# --- Tesseract Installation (Required for reading images) ---
# Pytesseract requires Google's Tesseract-OCR engine to be installed on your system.
# 1. Windows: Download and run the installer from: https://github.com/UB-Mannheim/tesseract/wiki
#    During installation, make sure to note the installation path.
# 2. MacOS: `brew install tesseract`
# 3. Linux (Ubuntu/Debian): `sudo apt-get install tesseract-ocr`
#
# After installing, you might need to tell pytesseract where to find the executable.
# Uncomment the line below and replace the path with your Tesseract installation path if you get an error.
# For example, on Windows, it might be: r'C:\Program Files\Tesseract-OCR\tesseract.exe'
# import pytesseract
# pytesseract.pytesseract.tesseract_cmd = r'<path_to_your_tesseract_executable>'


# --- Library Imports ---
# We import everything here to ensure names are defined.
try:
    # NOTE: Langchain and vector store libraries are removed as per user request.
    from langchain_core.documents import Document
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
    LIBRARIES_AVAILABLE = True
except ImportError as e:
    st.error(f"A required library is not installed. Please check the installation instructions. Error: {e}")
    LIBRARIES_AVAILABLE = False


# --- DeepSeek API Configuration ---
DEEPSEEK_API_URL = "https://api.deepseek.com/chat/completions"

# --- Helper Functions ---

def load_documents_from_folder(folder_path: str) -> List[Document]:
    """
    Reads all supported files (.txt, .pdf, .docx, .png, .jpg) from a specified folder,
    extracts the text, and returns them as a list of Document objects.
    The content is now sent directly to the LLM instead of a vector store.

    Args:
        folder_path (str): The path to the folder containing the files.

    Returns:
        List[Document]: A list of documents with their content.
    """
    documents = []
    if not os.path.isdir(folder_path):
        st.error(f"The provided path '{folder_path}' is not a valid directory.")
        return []

    filenames = os.listdir(folder_path)
    if not filenames:
        st.warning(f"No files found in the directory: {folder_path}")
        return []

    # Display the found files in the sidebar
    with st.sidebar:
        st.info(f"Found {len(filenames)} files: {', '.join(filenames)}")

    for file_name in filenames:
        file_path = os.path.join(folder_path, file_name)
        text = ""
        try:
            # Handle .pdf files
            if file_name.endswith(".pdf"):
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() or ""
            # Handle .docx files
            elif file_name.endswith(".docx"):
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            # Handle .txt files
            elif file_name.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            # Handle image files (.png, .jpg, .jpeg) using OCR
            elif file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
                try:
                    text = pytesseract.image_to_string(Image.open(file_path))
                except pytesseract.TesseractNotFoundError:
                    st.error("Tesseract is not installed or not in your PATH. Please check instructions.")
                    continue
                except Exception as ocr_e:
                    st.error(f"OCR failed for {file_name}: {ocr_e}")
                    continue
            # If text was extracted, create a Document object
            if text:
                documents.append(Document(page_content=text, metadata={"source": file_name}))
        except Exception as e:
            st.error(f"Failed to read or process {file_name}: {e}")
    return documents


def get_deepseek_response(prompt: str, api_key: str, chat_history: List[Dict[str, str]]) -> Generator[str, None, None]:
    """
    Sends a request to the DeepSeek API and streams the response.
    """
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    messages = chat_history + [{"role": "user", "content": prompt}]
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "stream": True
    }

    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload, stream=True)
        response.raise_for_status()
        for chunk in response.iter_lines():
            if chunk:
                decoded_chunk = chunk.decode('utf-8')
                if decoded_chunk.startswith('data: '):
                    try:
                        json_data = json.loads(decoded_chunk[6:])
                        if 'choices' in json_data and json_data['choices']:
                            content = json_data['choices'][0].get('delta', {}).get('content', '')
                            if content:
                                yield content
                    except json.JSONDecodeError:
                        continue
    except requests.exceptions.RequestException as e:
        error_message = f"API request failed: {e}"
        st.error(error_message)
        yield error_message
    except Exception as e:
        error_message = f"An unexpected error occurred during API call: {e}"
        st.error(error_message)
        yield error_message

# --- Custom CSS for Enhanced Design ---
def apply_custom_css():
    """Applies custom CSS to improve the application's design."""
    custom_css = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
        
        html, body, [class*="st-"] {
            font-family: 'Inter', sans-serif;
        }

        /* Main app background */
        .stApp {
            background-color: #111827; /* Dark blue-gray background */
            color: #F9FAFB;
        }

        /* Hide Streamlit's default header and footer */
        #MainMenu {display: none;}
        footer {display: none;}
        .stApp > header {
            background-color: transparent;
        }

        /* Chat message styling */
        [data-testid="chat-message-container"] {
            border-radius: 12px;
            padding: 16px;
            margin-bottom: 12px;
            max-width: 85%;
        }

        /* Assistant's message bubble */
        [data-testid="chat-message-container"]:has(div[data-testid="stChatMessageContent"][aria-label="assistant's message"]) {
            background-color: #374151; /* Gray-700 */
        }
        
        /* User's message bubble */
        [data-testid="chat-message-container"]:has(div[data-testid="stChatMessageContent"][aria-label="user's message"]) {
            background-color: #1E40AF; /* Blue-800 */
        }
        
        /* Style the chat input box */
        [data-testid="stChatInput"] {
            background-color: #1F2937; /* Gray-800 */
            border-radius: 12px;
            border: 1px solid #4B5563; /* Gray-600 */
        }
        [data-testid="stChatInput"] textarea {
            color: #F9FAFB;
        }

        /* Sidebar styling */
        [data-testid="stSidebar"] {
            background-color: #1F2937; /* Gray-800 */
            border-right: 1px solid #374151;
        }
    </style>
    """
    st.markdown(custom_css, unsafe_allow_html=True)


# --- Streamlit App UI ---
def main():
    """The main function that runs the Streamlit application."""
    st.set_page_config(
        page_title="Legal Chatbot (DeepSeek Direct)", 
        layout="wide",
        initial_sidebar_state="collapsed" # Sidebar is hidden by default
    )
    
    apply_custom_css()

    st.title("⚖️ Legal Document Chatbot with DeepSeek")
    
    if not LIBRARIES_AVAILABLE:
        st.stop()
    if not API_KEY or "sk-" not in API_KEY:
        st.error("Your DeepSeek API Key is not set correctly at the top of the script.")
        st.stop()

    # --- Sidebar for Configuration ---
    with st.sidebar:
        st.header("Configuration")
        st.write(f"**Document Folder:** `{FOLDER_PATH}`")
        st.write(f"**API Key:** `{API_KEY[:6]}...{API_KEY[-4:]}`")

    # --- Automatic Document Loading ---
    if "docs_loaded" not in st.session_state:
        with st.spinner("Loading documents automatically..."):
            docs = load_documents_from_folder(FOLDER_PATH)
            if docs:
                st.session_state.loaded_docs = docs
                st.session_state.docs_loaded = True
                st.toast(f"Successfully loaded {len(docs)} documents.", icon="✅")
            else:
                st.session_state.loaded_docs = []
                st.session_state.docs_loaded = False
                st.warning("No documents were found or loaded on startup.")
    
    # --- Main Chat Interface ---
    if "messages" not in st.session_state:
        st.session_state.messages = []

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("Ask a question about your documents..."):
        if not st.session_state.get("docs_loaded", False):
            st.warning("Documents were not loaded. Please check the folder and restart.")
            st.stop()

        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Sending documents and question to AI..."):
                context_text = "\n\n---\n\n".join([f"Source: {doc.metadata['source']}\n\nContent: {doc.page_content}" for doc in st.session_state.loaded_docs])
                
                if not context_text:
                    st.warning("The loaded documents appear to be empty. Cannot answer the question.")
                    st.stop()

                prompt_template = f"""
                You are a highly specialized legal assistant. Your task is to answer the user's question based *only* on the provided context documents.
                Do not use any external knowledge. If the answer is not found in the context, state that clearly.
                After providing the answer, you MUST cite the source file(s) you used from the context.

                CONTEXT DOCUMENTS:
                ---
                {context_text}
                ---

                USER'S QUESTION:
                {prompt}
                """

                response_placeholder = st.empty()
                full_response = ""
                
                chat_history_for_api = [
                    {"role": m["role"], "content": m["content"]}
                    for m in st.session_state.messages[:-1]
                ]

                for chunk in get_deepseek_response(prompt_template, API_KEY, chat_history_for_api):
                    full_response += chunk
                    response_placeholder.markdown(full_response + "▌")
                response_placeholder.markdown(full_response)

        st.session_state.messages.append({"role": "assistant", "content": full_response})

if __name__ == "__main__":
    main()
